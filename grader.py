# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2026 Travis Fraser
"""
Exam Anonymizer & Grader
========================
Privacy-first local app: student names NEVER sent to Claude.
Only anonymous IDs + exam text + rubric are transmitted.
"""

import base64
import concurrent.futures
import copy
import csv
import difflib
import io
import json
import logging
import os
import re
import secrets
import sqlite3
import statistics as _stats
import threading
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urlparse

from dotenv import load_dotenv
load_dotenv()

import anthropic
import httpx
import docx
import ollama
import pdfplumber
import pypdfium2 as pdfium
from PIL import Image, ImageEnhance
from pypdf import PdfReader, PdfWriter
from flask import (Flask, flash, g, jsonify, redirect, render_template,
                   request, send_file, session, url_for)
from scoring import (letter_grade, clean_feedback,
                     consolidate_and_clean, finalize_scores)
try:
    from audit_grader import (start_audit, get_audit_status, get_cumulative_audit_stats,
                              AUDIT_DIR, _audit_abort)
    from generate_audit_report import load_audit_data, compute_aggregate, build_report
    _HAS_AUDIT = True
except ImportError:
    _HAS_AUDIT = False
    AUDIT_DIR = Path(__file__).parent / "data" / "audit_results"

try:
    from rubric_builder import (
        _new_session, get_session, delete_session, extract_questions,
        start_refinement, respond_to_refinement, skip_question,
        advance_question, jump_to_question, auto_enhance,
        build_enhanced_rubric, format_for_grading_prompt,
        generate_mapping, apply_mapping,
        save_draft, load_draft, get_draft_info, discard_draft,
    )
    _HAS_BUILDER = True
except ImportError:
    _HAS_BUILDER = False

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
UPLOAD_DIR = DATA_DIR / "uploads"
RUBRIC_DIR = DATA_DIR / "rubrics"
DB_PATH = DATA_DIR / "exam_grader.db"
ALLOWED_EXT = {".pdf"}
MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB

# Local vision model used for cover-page OCR (runs via Ollama — no data leaves machine).
# Swap for any Ollama vision model: "llava:13b", "minicpm-v", etc.
OLLAMA_VISION_MODEL = "llama3.2-vision"

# Claude model used for grading. Update here when upgrading models.
CLAUDE_MODEL = "claude-sonnet-4-6"

# Boundary re-grade: exams scoring within this margin of a letter grade
# threshold get automatically re-graded for verification (ETS best practice).
_GRADE_BOUNDARIES = [90, 80, 70, 60]
_BOUNDARY_MARGIN = 1.5  # percentage points

# Vague feedback detector: patterns that are too generic to help students learn.
# Research basis: Nazaretsky et al. 2026 (JCAL) — vague AI feedback is discounted
# by students and fails to produce learning gains.
_VAGUE_FEEDBACK = re.compile(
    r'^(good\s*(work|job|answer)|mostly correct|needs? improvement|'
    r'well done|incorrect\.?|correct\.?|partial credit(?: awarded)?|'
    r'see rubric|ok(ay)?|fair|poor|excellent|satisfactory)\.?$',
    re.IGNORECASE,
)

# Score/feedback contradiction detector patterns.
# Over-scored: feedback says deduction but score gives full credit.
# Under-scored: feedback says correct but score is very low.
_DEDUCTION_LANGUAGE = re.compile(
    r'\b(0 points?|zero points?|no (?:credit|points?|marks?)|'
    r'incorrect|wrong answer|not correct|missed|'
    r'did not|failed to|unable to|omitted|blank|'
    r'deduct|minus|lost|penalty|error in)\b',
    re.IGNORECASE,
)
_FULL_CREDIT_LANGUAGE = re.compile(
    r'\b(correct(?:ly)?|full (?:marks?|credit|points?)|'
    r'perfect|well done|excellent work|earned all)\b',
    re.IGNORECASE,
)
_HANDWRITING_UNCERTAINTY = re.compile(
    r'\b(appears to (?:say|read|write|show)|seems? to (?:say|read)|'
    r'hard to (?:read|decipher|make out)|illegib|unread|'
    r'possibly|unclear (?:hand)?writ|smudg|faint (?:writing|text|answer)|'
    r'cannot (?:clearly |fully )?(?:read|determine|make out)|'
    r'difficult to (?:read|interpret)|partially (?:visible|legible)|'
    r'best guess|interpret(?:ed|ing) as)\b',
    re.IGNORECASE,
)

# Shared Anthropic client — thread-safe, instantiated once.
_anthropic_client = anthropic.Anthropic(
    api_key=os.environ.get("ANTHROPIC_API_KEY"),
    timeout=httpx.Timeout(300.0, connect=10.0),
)

DATA_DIR.mkdir(exist_ok=True)
UPLOAD_DIR.mkdir(exist_ok=True)
RUBRIC_DIR.mkdir(exist_ok=True)

# Persistent file logger — survives app restarts
_log = logging.getLogger("rubrica")
_log.setLevel(logging.INFO)
_log_handler = logging.FileHandler(DATA_DIR / "grading.log", encoding="utf-8")
_log_handler.setFormatter(logging.Formatter("%(asctime)s  %(levelname)-7s  %(message)s"))
_log.addHandler(_log_handler)

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY") or secrets.token_hex(32)


@app.context_processor
def inject_private_mode():
    return {"private_mode": session.get("private_mode", False)}


@app.context_processor
def inject_dark_mode():
    return {"dark_mode": session.get("dark_mode", False)}


@app.context_processor
def inject_has_builder():
    return {"has_builder": _HAS_BUILDER}


@app.context_processor
def inject_has_audit():
    return {"has_audit": _HAS_AUDIT}


def _safe_redirect_back():
    """Return the referrer URL only if it's same-origin; otherwise fall back to /."""
    ref = request.form.get("next") or request.referrer
    if ref:
        parsed = urlparse(ref)
        if not parsed.netloc or parsed.netloc == request.host:
            return ref
    return "/"


@app.route("/toggle-dark-mode", methods=["POST"])
def toggle_dark_mode():
    session["dark_mode"] = not session.get("dark_mode", False)
    return redirect(_safe_redirect_back())


@app.context_processor
def inject_active_review():
    review_files = sorted(DATA_DIR.glob("review_*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not review_files:
        return {"active_review_id": None, "active_review_ocr_running": False}
    review_path = review_files[0]
    review_id   = review_path.stem[len("review_"):]
    with _ocr_jobs_lock:
        ocr_running = _ocr_jobs.get(review_id, {}).get("running", False)
    return {"active_review_id": review_id, "active_review_ocr_running": ocr_running}


@app.route("/toggle-private-mode", methods=["POST"])
def toggle_private_mode():
    session["private_mode"] = not session.get("private_mode", False)
    return redirect(_safe_redirect_back())

# ---------------------------------------------------------------------------
# Background grading job state
# ---------------------------------------------------------------------------
_grade_job: dict   = {"running": False, "total": 0, "done": 0, "failed": 0, "errors": []}
_grade_lock        = threading.Lock()
_grade_abort       = threading.Event()
GRADE_WORKERS      = 5  # concurrent grading threads

# Background OCR job state (keyed by review_id)
_ocr_jobs: dict    = {}
_ocr_jobs_lock     = threading.Lock()

# ---------------------------------------------------------------------------
# Database helpers
# ---------------------------------------------------------------------------

def get_db():
    if "db" not in g:
        g.db = sqlite3.connect(DB_PATH, detect_types=sqlite3.PARSE_DECLTYPES)
        g.db.row_factory = sqlite3.Row
        g.db.execute("PRAGMA journal_mode=WAL")
    return g.db


@app.teardown_appcontext
def close_db(exc=None):
    db = g.pop("db", None)
    if db is not None:
        db.close()


def init_db():
    db = sqlite3.connect(DB_PATH)
    db.executescript("""
        CREATE TABLE IF NOT EXISTS rubrics (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            version    TEXT NOT NULL,
            content    TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS exams (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            anon_id      TEXT UNIQUE NOT NULL,
            student_name TEXT NOT NULL,
            version      TEXT NOT NULL,
            batch        INTEGER NOT NULL,
            file_path    TEXT NOT NULL,
            uploaded_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            ocr_text     TEXT,
            ocr_at       TIMESTAMP,
            graded_at    TIMESTAMP,
            grade_data   TEXT
        );

        CREATE TABLE IF NOT EXISTS roster (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            first_name TEXT NOT NULL,
            last_name  TEXT NOT NULL,
            sid        TEXT
        );
    """)
    # Migrate existing DB — add columns if not present
    migrations = [
        ("exams",   "ocr_text",         "TEXT"),
        ("exams",   "ocr_at",           "TIMESTAMP"),
        ("rubrics", "rubric_file_path", "TEXT"),
        ("exams",   "student_sid",      "TEXT"),
        ("rubrics", "total_points",     "REAL"),
        ("rubrics", "enhanced_rubric",  "TEXT"),
        ("exams",   "reviewed",         "INTEGER DEFAULT 0"),
        ("roster",  "email",            "TEXT"),
    ]
    _ALLOWED_TABLES = {"exams", "rubrics", "roster"}
    for table, col, typedef in migrations:
        if table not in _ALLOWED_TABLES:
            raise ValueError(f"Migration target not in allowlist: {table}")
        cols = [row[1] for row in db.execute(f"PRAGMA table_info({table})").fetchall()]
        if col not in cols:
            db.execute(f"ALTER TABLE {table} ADD COLUMN {col} {typedef}")
    db.commit()
    db.close()


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

_REVIEW_ID_RE = re.compile(r'^[0-9a-f]{16}$')


def _extract_json(text: str) -> str | None:
    """Return the first complete JSON object found in text, or None."""
    start = text.find("{")
    if start == -1:
        return None
    depth, in_string, escape = 0, False, False
    for i in range(start, len(text)):
        ch = text[i]
        if escape:
            escape = False
            continue
        if ch == '\\' and in_string:
            escape = True
            continue
        if ch == '"' and not escape:
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                return text[start:i + 1]
    return None


def _valid_review_id(review_id: str) -> bool:
    """Return True if review_id matches the expected hex token format."""
    return bool(_REVIEW_ID_RE.match(review_id))


def generate_anon_id():
    """8-char URL-safe random ID that won't be guessable or sequential."""
    while True:
        token = secrets.token_urlsafe(6)[:8].upper()
        db = get_db()
        exists = db.execute("SELECT 1 FROM exams WHERE anon_id=?", (token,)).fetchone()
        if not exists:
            return token


def extract_pdf_text(filepath: str) -> str:
    """Extract all text from a PDF. Returns empty string on failure."""
    try:
        with pdfplumber.open(filepath) as pdf:
            return "\n".join(
                page.extract_text() or "" for page in pdf.pages
            ).strip()
    except Exception as e:
        return f"[PDF extraction error: {e}]"


def extract_docx_text(filepath: str) -> str:
    """Extract all text from a .docx file, preserving paragraph breaks."""
    try:
        doc = docx.Document(filepath)
        return "\n".join(para.text for para in doc.paragraphs).strip()
    except Exception as e:
        return f"[DOCX extraction error: {e}]"


def extract_rubric_file(filepath: str, filename: str) -> str:
    """Extract rubric text from a PDF or DOCX file."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_pdf_text(filepath)
    elif ext in (".docx", ".doc"):
        return extract_docx_text(filepath)
    return ""


def get_rubric(version: str):
    db = get_db()
    return db.execute(
        "SELECT * FROM rubrics WHERE version=? ORDER BY id DESC LIMIT 1",
        (version,)
    ).fetchone()


def _is_boundary_score(pct: float) -> bool:
    """True if percentage falls within +/-1.5 of a letter grade threshold."""
    return any(abs(pct - b) <= _BOUNDARY_MARGIN for b in _GRADE_BOUNDARIES)


# ---------------------------------------------------------------------------
# Roster fuzzy matching — 100% local, zero network calls
# ---------------------------------------------------------------------------

def match_against_roster(extracted_name: str, extracted_sid: str, roster_entries: list) -> tuple:
    """
    Pure local fuzzy match against roster entries.
    Normalises common OCR digit/letter confusions in SID (O→0, l→1, I→1, S→5).
    Returns (best_entry_dict, score 0-1) or (None, 0.0).
    PRIVACY: roster data never leaves this function — no API calls made here.
    """
    if not roster_entries:
        return None, 0.0

    def norm_sid(s):
        return s.upper().replace('O', '0').replace('l', '1').replace('I', '1').replace('S', '5')

    SM = difflib.SequenceMatcher
    name_clean = (extracted_name or "").lower().strip()
    sid_clean  = (extracted_sid  or "").strip()
    best_score, best_entry = 0.0, None

    for entry in roster_entries:
        fl = f"{entry['first_name']} {entry['last_name']}".lower()
        lf = f"{entry['last_name']} {entry['first_name']}".lower()
        name_score = max(SM(None, name_clean, fl).ratio(),
                         SM(None, name_clean, lf).ratio())

        sid_score = 0.0
        if sid_clean and entry['sid']:
            a, b = norm_sid(sid_clean), norm_sid(entry['sid'])
            sid_score = 1.0 if a == b else SM(None, a, b).ratio()

        has_both = bool(name_clean) and bool(sid_clean and entry['sid'])
        if has_both:
            score = 0.55 * name_score + 0.45 * sid_score
        elif name_clean:
            score = name_score
        else:
            score = sid_score

        if score > best_score:
            best_score, best_entry = score, entry

    return best_entry, round(best_score, 3)


# ---------------------------------------------------------------------------
# OCR — render PDF pages as images and transcribe with Claude vision
# ---------------------------------------------------------------------------

_pdfium_lock = threading.Lock()

def _render_page_png(pdf_path: str, page_num: int, scale: float = 2.0) -> bytes:
    with _pdfium_lock:
        doc = pdfium.PdfDocument(pdf_path)
        try:
            page    = doc[page_num]
            bitmap  = page.render(scale=scale)
            pil_img = bitmap.to_pil()
            buf     = io.BytesIO()
            pil_img.save(buf, format="PNG")
            result  = buf.getvalue()
            # Explicitly release PDFium sub-objects before closing the document.
            # Python's GC may not free them promptly, which corrupts PDFium's
            # internal state across multiple calls (particularly on Windows).
            del pil_img, bitmap, page
            return result
        finally:
            doc.close()



def _cover_phash(pdf_path: str, hash_size: int = 8) -> int | None:
    """Render cover page at low resolution and return a perceptual hash integer."""
    try:
        img_bytes = _render_page_png(pdf_path, 0, scale=0.3)
        img = Image.open(io.BytesIO(img_bytes)).convert("L").resize(
            (hash_size, hash_size), Image.LANCZOS
        )
        pixels = list(img.getdata())
        avg    = sum(pixels) / len(pixels)
        return sum(1 << i for i, p in enumerate(pixels) if p > avg)
    except Exception:
        return None


def check_cover_consistency(exams: list) -> None:
    """
    Compare cover page perceptual hashes across all exams and flag outliers in-place.
    Adds 'cover_flag': True to any exam whose cover differs significantly from the
    majority. Runs entirely locally — no API calls. Modifies exams in-place.
    """
    if len(exams) < 2:
        for exam in exams:
            exam["cover_flag"] = False
        return

    THRESHOLD = 15  # Hamming distance out of 64 bits
    hashes = [_cover_phash(exam["file_path"]) for exam in exams]
    valid  = [(i, h) for i, h in enumerate(hashes) if h is not None]

    for i, exam in enumerate(exams):
        if hashes[i] is None:
            exam["cover_flag"] = False
            continue
        distances = [bin(hashes[i] ^ h).count("1") for j, h in valid if j != i]
        avg_dist  = sum(distances) / len(distances) if distances else 0
        exam["cover_flag"] = avg_dist > THRESHOLD


def read_name_sid_from_cover(file_path: str) -> dict:
    """
    Extract student name and SID from cover page (page 0) using a local Ollama
    vision model (OLLAMA_VISION_MODEL).  Runs entirely on-device via GPU —
    no data leaves the machine.
    Returns {"name": "...", "sid": "..."}.  Falls back to empty strings on any failure.
    """
    try:
        img_bytes = _render_page_png(file_path, 0, scale=1.5)
        response  = ollama.chat(
            model=OLLAMA_VISION_MODEL,
            messages=[{
                "role": "user",
                "content": (
                    "Look at this exam cover page. Extract the student's full name "
                    "and student ID number (labelled SID, ID, Student Number, or similar). "
                    "Respond with ONLY valid JSON, exactly: "
                    '{"name": "Full Name", "sid": "123456789"} '
                    "Use an empty string if you cannot find the value. "
                    "No extra text, no markdown fences — just the JSON object."
                ),
                "images": [img_bytes],
            }],
        )
        text = response["message"]["content"].strip()
        # Strip markdown code fences if present
        if text.startswith("```"):
            lines = text.splitlines()
            text  = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
        data = json.loads(text)
        return {"name": str(data.get("name", "")).strip(),
                "sid":  str(data.get("sid",  "")).strip()}
    except Exception as e:
        _log.warning("Cover OCR failed for %s: %s", file_path, e)
        return {"name": "", "sid": ""}


def _run_ocr_background(review_id: str, review_path: Path):
    """Run cover-page OCR for every exam in a review session on a background thread.
    After OCR completes, auto-matches against the roster if one is loaded."""
    try:
        data  = json.loads(review_path.read_text())
        total = len(data["exams"])
        with _ocr_jobs_lock:
            _ocr_jobs[review_id] = {
                "total": total, "done": 0, "running": True, "aborted": False,
                "auto_matched": False,
            }
        for i, exam in enumerate(data["exams"]):
            with _ocr_jobs_lock:
                if _ocr_jobs[review_id].get("aborted"):
                    break
            result       = read_name_sid_from_cover(exam["file_path"])
            exam["name"] = result["name"]
            exam["sid"]  = result["sid"]
            tmp = review_path.with_suffix(".tmp")
            tmp.write_text(json.dumps(data))
            tmp.replace(review_path)
            with _ocr_jobs_lock:
                _ocr_jobs[review_id]["done"] = i + 1

        # --- Auto-match against roster after OCR ---
        with _ocr_jobs_lock:
            aborted = _ocr_jobs.get(review_id, {}).get("aborted", False)
        if not aborted:
            try:
                with app.app_context():
                    db = get_db()
                    roster_count = db.execute("SELECT COUNT(*) FROM roster").fetchone()[0]
                    if roster_count > 0:
                        rows = db.execute("SELECT * FROM roster").fetchall()
                        entries = [dict(r) for r in rows]
                        # Re-read review data (OCR may have updated it)
                        data = json.loads(review_path.read_text())
                        matched = 0
                        for exam in data["exams"]:
                            best, score = match_against_roster(
                                exam.get("name", ""), exam.get("sid", ""), entries
                            )
                            if best and score >= 0.60:
                                exam["name"] = f"{best['first_name']} {best['last_name']}"
                                exam["sid"] = best["sid"] or ""
                                exam["roster_score"] = round(score, 3)
                                matched += 1
                        tmp = review_path.with_suffix(".tmp")
                        tmp.write_text(json.dumps(data))
                        tmp.replace(review_path)
                        with _ocr_jobs_lock:
                            _ocr_jobs[review_id]["auto_matched"] = True
                        _log.info("[ROSTER] Auto-matched %d/%d exams for review %s",
                                  matched, len(data["exams"]), review_id)
            except Exception as e:
                _log.warning("[ROSTER] Auto-match failed for review %s: %s", review_id, e)

    except Exception as e:
        _log.error("OCR failed for review %s: %s", review_id, e)
        with _ocr_jobs_lock:
            if review_id in _ocr_jobs:
                _ocr_jobs[review_id]["error"] = str(e)
    finally:
        with _ocr_jobs_lock:
            if review_id in _ocr_jobs:
                _ocr_jobs[review_id]["running"] = False


# ---------------------------------------------------------------------------
# Batch PDF splitting — fixed page count
# ---------------------------------------------------------------------------

def split_by_page_count_bytes(pdf_bytes: bytes, pages_per_exam: int,
                              output_dir: Path) -> list[dict]:
    """
    Split a PDF (supplied as raw bytes) into chunks of `pages_per_exam` pages.
    Works entirely in memory — no temp file is ever opened, avoiding Windows
    file-locking issues.
    Returns list of dicts: {name, file_path, pages}.
    """
    reader = PdfReader(io.BytesIO(pdf_bytes))
    total  = len(reader.pages)
    exams  = []

    for start in range(0, total, pages_per_exam):
        end = min(start + pages_per_exam - 1, total - 1)

        writer = PdfWriter()
        for p in range(start, end + 1):
            writer.add_page(reader.pages[p])

        buf = io.BytesIO()
        writer.write(buf)

        temp_id  = secrets.token_hex(8)
        out_path = output_dir / f"tmp_{temp_id}.pdf"
        out_path.write_bytes(buf.getvalue())

        exams.append({
            "name":       "",
            "sid":        "",
            "file_path":  str(out_path),
            "pages":      f"{start + 1}–{end + 1}",
            "page_count": end - start + 1,
        })

    return exams



# ---------------------------------------------------------------------------
# Routes: Dashboard
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    db = get_db()
    total   = db.execute("SELECT COUNT(*) FROM exams").fetchone()[0]
    graded  = db.execute("SELECT COUNT(*) FROM exams WHERE grade_data IS NOT NULL").fetchone()[0]
    rubrics = db.execute("SELECT DISTINCT version FROM rubrics").fetchall()
    versions_with_rubrics = [r["version"] for r in rubrics]
    roster_count = db.execute("SELECT COUNT(*) FROM roster").fetchone()[0]
    audit_stats = get_cumulative_audit_stats() if _HAS_AUDIT else {"total_audited": 0, "runs": 0}
    return render_template("index.html",
                           total=total,
                           graded=graded,
                           ungraded=total - graded,
                           versions_with_rubrics=versions_with_rubrics,
                           roster_count=roster_count,
                           audit_stats=audit_stats)


# ---------------------------------------------------------------------------
# Routes: Rubric Setup
# ---------------------------------------------------------------------------

@app.route("/setup", methods=["GET", "POST"])
def setup():
    db = get_db()
    if request.method == "POST":
        version      = request.form.get("version", "").strip().upper()
        file         = request.files.get("rubric_file")
        total_points = request.form.get("total_points", type=float)

        if not version:
            flash("Exam version is required.", "danger")
            return redirect(url_for("setup"))

        if not file or not file.filename:
            flash("Upload a rubric file (.pdf or .docx).", "danger")
            return redirect(url_for("setup"))

        ext = Path(file.filename).suffix.lower()
        if ext not in (".pdf", ".docx", ".doc"):
            flash("Only PDF or Word (.docx) files are supported.", "danger")
            return redirect(url_for("setup"))

        # Save permanently so we can send the original file to Claude when grading
        perm_path = RUBRIC_DIR / f"{version}_{secrets.token_hex(6)}{ext}"
        file.save(str(perm_path))
        rubric_file_path = str(perm_path)

        # Extract text for display / fallback
        content = extract_rubric_file(str(perm_path), file.filename)
        if not content or content.startswith("["):
            content = f"[See uploaded rubric file — text extraction unavailable for this file]"

        db.execute(
            "INSERT INTO rubrics (version, content, rubric_file_path, total_points) VALUES (?, ?, ?, ?)",
            (version, content, rubric_file_path, total_points)
        )
        db.commit()
        flash(f"Rubric for Version {version} saved.", "success")
        return redirect(url_for("setup"))

    rubrics = db.execute(
        "SELECT DISTINCT version, MAX(created_at) as updated, enhanced_rubric FROM rubrics GROUP BY version"
    ).fetchall()
    section = request.args.get("section", "upload")
    builder_rubrics = []
    draft_info = None
    if _HAS_BUILDER and section == "builder":
        builder_rubrics = db.execute(
            "SELECT DISTINCT version, MAX(created_at) as updated, total_points, enhanced_rubric "
            "FROM rubrics GROUP BY version"
        ).fetchall()
        draft_info = get_draft_info()
    # Find the primary enhanced version (has enhanced_rubric without mapped_from)
    primary_enhanced = None
    for br in (builder_rubrics if builder_rubrics else rubrics):
        if br["enhanced_rubric"]:
            try:
                edata = json.loads(br["enhanced_rubric"])
                if not edata.get("mapped_from"):
                    primary_enhanced = br["version"]
                    break
            except (json.JSONDecodeError, TypeError):
                pass
    return render_template("setup.html", rubrics=rubrics, section=section,
                           builder_rubrics=builder_rubrics, draft_info=draft_info,
                           primary_enhanced=primary_enhanced)


@app.route("/setup/delete/<version>", methods=["POST"])
def delete_rubric(version):
    db = get_db()
    db.execute("DELETE FROM rubrics WHERE version=?", (version.upper(),))
    db.commit()
    flash(f"Rubric for Version {version.upper()} deleted.", "success")
    return redirect(url_for("setup"))


@app.route("/setup/view/<version>")
def view_rubric(version):
    rubric = get_rubric(version.upper())
    if not rubric:
        flash(f"No rubric found for Version {version}.", "warning")
        return redirect(url_for("setup"))
    return render_template("view_rubric.html", rubric=rubric)


# ---------------------------------------------------------------------------
# Routes: View Exams (anonymized — no student names shown by default)
# ---------------------------------------------------------------------------

@app.route("/exams")
def exams():
    db   = get_db()
    rows = db.execute(
        """SELECT anon_id, version, batch, uploaded_at,
                  CASE WHEN grade_data IS NOT NULL THEN 1 ELSE 0 END as graded
           FROM exams ORDER BY version, batch, uploaded_at"""
    ).fetchall()
    rubric_versions = {
        r["version"] for r in db.execute("SELECT DISTINCT version FROM rubrics").fetchall()
    }
    exam_versions = {r["version"] for r in rows}
    missing_rubric_versions = sorted(exam_versions - rubric_versions)
    return render_template("exams.html", exams=rows,
                           rubric_versions=rubric_versions,
                           missing_rubric_versions=missing_rubric_versions)


# ---------------------------------------------------------------------------
# Routes: Delete exam
# ---------------------------------------------------------------------------

@app.route("/exams/delete-all", methods=["POST"])
def delete_all_exams():
    db   = get_db()
    rows = db.execute("SELECT file_path FROM exams").fetchall()
    deleted_files, failed_files = 0, 0
    for row in rows:
        try:
            Path(row["file_path"]).unlink(missing_ok=True)
            deleted_files += 1
        except Exception as e:
            _log.warning(f"Could not delete file {row['file_path']}: {e}")
            failed_files += 1
    db.execute("DELETE FROM exams")
    db.commit()
    msg = f"All exams removed ({deleted_files} file{'s' if deleted_files != 1 else ''} deleted"
    if failed_files:
        msg += f", {failed_files} file(s) could not be removed from disk"
    flash(msg + ".", "success" if not failed_files else "warning")
    return redirect(url_for("exams"))


@app.route("/exam/<anon_id>/delete", methods=["POST"])
def delete_exam(anon_id):
    db   = get_db()
    exam = db.execute("SELECT file_path FROM exams WHERE anon_id=?", (anon_id,)).fetchone()
    if not exam:
        flash("Exam not found.", "danger")
        return redirect(url_for("exams"))

    # Delete the PDF file from disk
    try:
        Path(exam["file_path"]).unlink(missing_ok=True)
    except Exception as e:
        _log.warning(f"Could not delete file for {anon_id}: {e}")

    db.execute("DELETE FROM exams WHERE anon_id=?", (anon_id,))
    db.commit()
    flash(f"Exam {anon_id} removed.", "success")
    return redirect(url_for("exams"))


@app.route("/delete-selected", methods=["POST"])
def delete_selected():
    anon_ids = request.form.getlist("anon_ids")
    if not anon_ids:
        return jsonify({"error": "No exams selected."}), 400
    db = get_db()
    removed, failed = 0, 0
    for aid in anon_ids:
        exam = db.execute("SELECT file_path FROM exams WHERE anon_id=?", (aid,)).fetchone()
        if not exam:
            continue
        try:
            Path(exam["file_path"]).unlink(missing_ok=True)
        except Exception as e:
            _log.warning(f"Could not delete file for {aid}: {e}")
            failed += 1
        db.execute("DELETE FROM exams WHERE anon_id=?", (aid,))
        removed += 1
    db.commit()
    return jsonify({"removed": removed, "failed": failed})


# ---------------------------------------------------------------------------
# Routes: OCR
# ---------------------------------------------------------------------------



# ---------------------------------------------------------------------------
# Routes: Grading
# ---------------------------------------------------------------------------

def _grade_exam(exam_row, rubric) -> dict:
    """
    Grade an exam by sending page images directly to Claude Sonnet (vision grading).
    Page 0 (name/cover sheet) is skipped — only answer pages are sent.
    If the rubric was uploaded as a PDF, it is sent natively (preserving tables/charts).
    PRIVACY: only anon_id, version, and page images are sent. No student name.
    """
    doc   = pdfium.PdfDocument(exam_row["file_path"])
    total = len(doc)
    doc.close()
    if total < 2:
        raise ValueError("Exam PDF must have at least 2 pages (cover + answers).")

    # --- Rubric block: send as native PDF if available, otherwise as text ---
    rubric_file = rubric["rubric_file_path"] if rubric["rubric_file_path"] else None
    rubric_blocks = []
    if rubric_file and Path(rubric_file).exists() and rubric_file.lower().endswith(".pdf"):
        with open(rubric_file, "rb") as f:
            rubric_b64 = base64.standard_b64encode(f.read()).decode()
        rubric_blocks = [
            {"type": "text", "text": "RUBRIC (read the document below carefully before grading):"},
            {
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": rubric_b64},
            },
        ]
    else:
        # Typed rubric or DOCX (text extraction fallback)
        rubric_blocks = [
            {"type": "text", "text": f"RUBRIC:\n{rubric['content']}"},
        ]

    # --- Exam answer page images (skip page 0 — name/cover sheet) ---
    exam_blocks = []
    for page_num in range(1, total):
        img_bytes = _render_page_png(exam_row["file_path"], page_num)
        img = Image.open(io.BytesIO(img_bytes))
        img = ImageEnhance.Contrast(img).enhance(2.0)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        img_bytes = buf.getvalue()
        b64       = base64.standard_b64encode(img_bytes).decode()
        exam_blocks.append({
            "type": "image",
            "source": {"type": "base64", "media_type": "image/png", "data": b64},
        })
        exam_blocks.append({"type": "text", "text": f"[Page {page_num}]"})

    # --- Enhanced rubric (partial credit guide) if available ---
    enhanced_blocks = []
    if _HAS_BUILDER and rubric["enhanced_rubric"]:
        try:
            enhanced = json.loads(rubric["enhanced_rubric"])
            enhanced_text = format_for_grading_prompt(enhanced)
            enhanced_blocks = [{"type": "text", "text": enhanced_text}]
        except (json.JSONDecodeError, TypeError):
            pass

    # Version-correct rubric text for refinement/contradiction sub-calls.
    # rubric["content"] may contain the primary version's text (wrong Qs for mapped versions).
    # The enhanced rubric text has the correct question numbers for this version.
    if _HAS_BUILDER and rubric["enhanced_rubric"]:
        try:
            _enhanced_for_text = json.loads(rubric["enhanced_rubric"])
            rubric_text_for_subcalls = format_for_grading_prompt(_enhanced_for_text)
        except (json.JSONDecodeError, TypeError):
            rubric_text_for_subcalls = rubric["content"] or "(rubric not available)"
    else:
        rubric_text_for_subcalls = rubric["content"] or "(rubric not available)"

    content = rubric_blocks + enhanced_blocks + [
        {"type": "text", "text": "--- EXAM ANSWER PAGES ---"},
        *exam_blocks,
        {
            "type": "text",
            "text": (
                f"Please grade this exam.\n\n"
                f"EXAM ID: {exam_row['anon_id']}\n"
                f"VERSION: {exam_row['version']}\n\n"
                f"Grade each question strictly according to the rubric above. "
                f"Respond ONLY with the JSON object specified in the system prompt."
            ),
        },
    ]

    system_prompt = f"""You are an impartial exam grader. You will grade handwritten exams anonymously by reading the page images provided.

EXAM VERSION: {exam_row["version"]}
ANONYMOUS EXAM ID: {exam_row["anon_id"]}

Grading instructions:
- The rubric is provided at the start of the user message.
- Read the handwritten answers directly from the exam page images.
- Grade each question strictly according to the rubric.
- Some answers may be written in faint pencil — examine the image carefully before concluding a question is blank or unanswered.
- Be fair and consistent. Do not infer the student's identity.

MULTIPLE CHOICE instructions (critical):
- Each MC question lists options VERTICALLY in order: A (top), B (second), C (third), D (bottom).
- The student circles exactly one letter. That circled letter is their answer.
- READ THE ACTUAL LETTER INSIDE OR NEXT TO THE CIRCLE — do not infer the answer from position alone.
- B and D look similar when handwritten. Look carefully: B has two bumps on the right; D has one large curve.
- If a circle is around the second option, confirm it says "B" not "D" before recording.
- The circled letter IS the student's answer — compare it directly to the correct answer in the rubric table.
- Award full points if the circled letter matches the correct answer. Award 0 otherwise.
- Do NOT award partial credit on multiple choice.
- If no letter is clearly marked, award 0 and note "no answer marked" in feedback.
- If two letters appear marked, use the one with the clearest, most deliberate marking.
- CROSSED-OUT ANSWERS: Students often cross out wrong answers with slashes (/) or Xs through the letter or circle. A crossed-out/slashed answer is REJECTED by the student — ignore it completely. Only score the answer that is cleanly circled without cross-out marks. If ALL circled answers are crossed out, treat as "no answer marked" (0 pts).

QUESTION CONSOLIDATION (critical):
- Each question must appear as a SINGLE entry in the scores array, even if the rubric breaks it into sub-parts (a), (b), (c).
- Sum all sub-part points into one earned_points and one max_points for the parent question.
- Use the question number only (e.g. "Q3", "Q22") — never "Q3a", "Q3b", "Q22a", etc.
- Include feedback for all sub-parts combined in the single feedback string.

FEEDBACK TONE (critical):
- The feedback field is shown DIRECTLY TO STUDENTS. It must read as a final verdict, not a grading worksheet.
- Write as a professor would on a graded exam — definitive, concise, and authoritative.
- State what is correct or incorrect directly. Never hedge, self-correct, or show reasoning process.
- NEVER include score tallies, recalculations, point breakdowns, or corrections in feedback. The scores are computed separately — feedback explains WHY, not HOW MANY.
- NEVER use phrases like "wait", "actually", "let me reconsider", "on second thought", "correcting", "awarding X pts", "total = X", "this earns", or similar deliberation/arithmetic language.
- For multi-part questions, address each part's correctness without re-deriving the total.
- If you are uncertain about a reading, make your best judgment and commit to it. Do not narrate your uncertainty.
- Good: "Correct application of the Coase theorem."
- Good: "Part (a): Correctly identifies the fairness effect. Part (b): Movement along the curve is correct; shift identified with valid non-price factor."
- Bad: "Part (a) earns 1 pt. Part (b) earns 1.5 pts. Total = 2.5. Correcting: actually awarding 3 pts..."

HANDWRITING UNCERTAINTY (required for every question):
- You MUST include "handwriting_flag" (true or false) in EVERY question object.
- Set true if: the handwriting is ambiguous, partially illegible, you guessed between
  two possible letters/digits, the writing is smudged or overlapping, or you made any
  judgment call about what was written.
- Set false if: the handwriting is clearly legible with no ambiguity.
- Omitting this field is an error. Every question MUST have it.

- Respond ONLY with valid JSON in exactly this format:
{{
  "anon_id": "{exam_row["anon_id"]}",
  "scores": [
    {{"question": "Q1", "max_points": <n>, "earned_points": <n>, "feedback": "<specific feedback>", "handwriting_flag": false}}
  ],
  "total_earned": <n>,
  "total_possible": <n>,
  "letter_grade": "<A/B/C/D/F>",
  "overall_feedback": "<2-3 sentence summary>"
}}"""

    client = _anthropic_client

    MAX_ATTEMPTS = 2
    last_error = None
    data = None

    for attempt in range(1, MAX_ATTEMPTS + 1):
        with client.messages.stream(
            model=CLAUDE_MODEL,
            max_tokens=16384,
            temperature=0.0,
            system=system_prompt,
            messages=[{"role": "user", "content": content}],
        ) as stream:
            full_response = stream.get_final_message()

        if full_response.stop_reason == "max_tokens":
            last_error = ValueError("Response was cut off (max_tokens reached). The exam may be too long.")
            _log.warning("Attempt %d/%d for %s: max_tokens reached (limit=16384)", attempt, MAX_ATTEMPTS, exam_row["anon_id"])
            continue

        response_text = ""
        for block in full_response.content:
            if block.type == "text":
                response_text = block.text
                break

        json_str = _extract_json(response_text)
        if json_str is None:
            last_error = ValueError(f"No JSON found. Claude responded with: {response_text[:500]}")
            _log.warning("Attempt %d/%d for %s: no JSON in response", attempt, MAX_ATTEMPTS, exam_row["anon_id"])
            continue

        try:
            data = json.loads(json_str)
            break  # success
        except json.JSONDecodeError as e:
            last_error = e
            _log.warning("Attempt %d/%d for %s: malformed JSON - %s", attempt, MAX_ATTEMPTS, exam_row["anon_id"], e)
            continue

    if data is None:
        raise last_error or ValueError("Grading failed: no valid response after retries")

    # Phase 1: sub-part consolidation + feedback sanitization (shared with audit)
    consolidate_and_clean(data)

    # --- Raw score snapshot (for audit comparison) ---
    # Capture model output after shared pipeline (consolidation + sanitization)
    # but before production-only safeguards (feedback specificity, contradiction
    # resolution, boundary re-grade). finalize_scores() on the copy produces the
    # same state audit_grader.py would, making comparison apples-to-apples.
    raw_snapshot = copy.deepcopy({"scores": data.get("scores", []),
                                  "total_earned": data.get("total_earned", 0),
                                  "total_possible": data.get("total_possible", 0)})
    finalize_scores(raw_snapshot)
    data["raw_scores"] = raw_snapshot

    # --- Feedback specificity enforcement ---
    # Vague feedback like "Good work" or "Incorrect" without explanation is
    # pedagogically useless (Nazaretsky et al. 2026). Detect and refine via
    # a lightweight text-only API call (no images, low token cost).
    # Skip full-marks questions — "Correct" is sufficient for MC and acceptable
    # for free-response when nothing was missed. Detailed feedback matters most
    # when points are lost.
    vague_questions = []
    if data.get("scores"):
        for s in data["scores"]:
            # Full marks → skip (MC correct, perfect free-response — nothing to explain)
            if s.get("earned_points", 0) >= s.get("max_points", 1):
                continue
            fb = (s.get("feedback") or "").strip()
            if not fb or len(fb.split()) < 8 or _VAGUE_FEEDBACK.match(fb):
                vague_questions.append(s)

    if vague_questions:
        _log.info("[FEEDBACK] %s has %d question(s) with vague feedback - requesting refinement",
                  exam_row["anon_id"], len(vague_questions))
        q_details = "\n".join(
            f"- {q['question']}: earned {q['earned_points']}/{q['max_points']}. "
            f"Current feedback: \"{q.get('feedback', '')}\""
            for q in vague_questions
        )
        rubric_text = rubric_text_for_subcalls
        refine_prompt = (
            f"You graded an exam and gave feedback that is too vague to help the student.\n\n"
            f"RUBRIC:\n{rubric_text}\n\n"
            f"Questions needing better feedback:\n{q_details}\n\n"
            f"For each question, write specific feedback (15+ words) that states what "
            f"was correct or incorrect and references the rubric criteria.\n\n"
            f"Respond ONLY with JSON: {{\"improved\": [{{\"question\": \"Q1\", \"feedback\": \"...\"}}]}}"
        )
        try:
            refine_resp = client.messages.create(
                model=CLAUDE_MODEL,
                max_tokens=2048,
                temperature=0.0,
                messages=[{"role": "user", "content": refine_prompt}],
            )
            refine_text = ""
            for block in refine_resp.content:
                if block.type == "text":
                    refine_text = block.text
                    break
            refine_json = _extract_json(refine_text)
            if refine_json:
                improved = json.loads(refine_json)
                fb_map = {item["question"]: item["feedback"]
                          for item in improved.get("improved", [])}
                for s in data["scores"]:
                    if s["question"] in fb_map and len(fb_map[s["question"]].split()) >= 10:
                        s["feedback"] = clean_feedback(fb_map[s["question"]])
                        _log.info("[FEEDBACK] Refined %s for %s",
                                  s["question"], exam_row["anon_id"])
        except Exception as e:
            _log.warning("[FEEDBACK] Refinement failed for %s: %s", exam_row["anon_id"], e)

    # --- Score/feedback semantic contradiction detection & resolution ---
    # Catches cases where feedback language contradicts the score direction.
    # Point tallies are stripped from feedback by clean_feedback() so we only
    # need to detect semantic mismatches (e.g. "incorrect" with full marks).
    contradiction_resolved = []
    contradictions = []
    if data.get("scores"):
        for s in data["scores"]:
            fb = (s.get("feedback") or "").strip()
            earned = float(s.get("earned_points", 0))
            mx = float(s.get("max_points", 1))
            has_deduction = bool(_DEDUCTION_LANGUAGE.search(fb))
            has_credit = bool(_FULL_CREDIT_LANGUAGE.search(fb))
            # Over-scored: feedback says deduction but score gives full credit.
            # Skip mixed feedback (both patterns present = consolidated sub-parts).
            if has_deduction and not has_credit and earned >= mx:
                contradictions.append(s)
            # Under-scored: feedback says correct but score is very low.
            elif has_credit and not has_deduction and earned < mx * 0.5:
                contradictions.append(s)

    if contradictions:
        _log.info("[CONTRADICTION] %s has %d question(s) with score/feedback contradictions",
                  exam_row["anon_id"], len(contradictions))
        rubric_text = rubric_text_for_subcalls
        c_details = "\n".join(
            f"- {c['question']}: earned {c['earned_points']}/{c['max_points']}. "
            f"Feedback: \"{c.get('feedback', '')}\""
            for c in contradictions
        )
        resolve_prompt = (
            f"You graded an exam but your feedback and scores contradict each other.\n\n"
            f"RUBRIC:\n{rubric_text}\n\n"
            f"Contradictions found:\n{c_details}\n\n"
            f"For each question, re-examine whether the score or the feedback is correct. "
            f"Return the corrected earned_points and feedback.\n\n"
            f"Respond ONLY with JSON: "
            f"{{\"resolved\": [{{\"question\": \"Q1\", \"earned_points\": <n>, \"feedback\": \"...\"}}]}}"
        )
        try:
            resolve_resp = client.messages.create(
                model=CLAUDE_MODEL,
                max_tokens=2048,
                temperature=0.0,
                messages=[{"role": "user", "content": resolve_prompt}],
            )
            resolve_text = ""
            for block in resolve_resp.content:
                if block.type == "text":
                    resolve_text = block.text
                    break
            resolve_json = _extract_json(resolve_text)
            if resolve_json:
                resolved = json.loads(resolve_json)
                fix_map = {item["question"]: item for item in resolved.get("resolved", [])}
                for s in data["scores"]:
                    if s["question"] in fix_map:
                        fix = fix_map[s["question"]]
                        new_earned = float(fix.get("earned_points", s["earned_points"]))
                        new_fb = fix.get("feedback", "").strip()
                        mx = float(s["max_points"])
                        # Apply only if earned is valid and feedback is substantive
                        if 0 <= new_earned <= mx and len(new_fb.split()) >= 10:
                            _log.info("[CONTRADICTION] %s %s: %s/%s -> %s/%s",
                                      exam_row["anon_id"], s["question"],
                                      s["earned_points"], s["max_points"],
                                      new_earned, s["max_points"])
                            s["earned_points"] = new_earned
                            s["feedback"] = clean_feedback(new_fb)
                            contradiction_resolved.append(s["question"])
        except Exception as e:
            _log.warning("[CONTRADICTION] Resolution failed for %s: %s", exam_row["anon_id"], e)

    # --- MC double-read verification ---
    # MC questions scored 0 are high-risk for letter-reading errors (3.33 pt swing).
    # Send a focused verification call for each MC zero to confirm the circled letter.
    _mc_letter_re = re.compile(
        r'\b(?:circled|selected|marked|chose|answered?)\s+([A-D])\b', re.IGNORECASE
    )
    mc_verified = []
    if data.get("scores"):
        for s in data["scores"]:
            earned = float(s.get("earned_points", 0))
            mx = float(s.get("max_points", 0))
            fb = s.get("feedback", "")
            page = s.get("page")
            # Detect MC zeros: all-or-nothing score, 0 earned, short feedback with letter ref
            if earned > 0 or mx <= 0 or page is None:
                continue
            letter_match = _mc_letter_re.search(fb)
            if not letter_match:
                continue
            original_letter = letter_match.group(1).upper()
            # Send just the relevant page for a focused letter-read
            try:
                verify_img = _render_page_png(exam_row["file_path"], page - 1)
                verify_img_obj = Image.open(io.BytesIO(verify_img))
                verify_img_obj = ImageEnhance.Contrast(verify_img_obj).enhance(2.0)
                vbuf = io.BytesIO()
                verify_img_obj.save(vbuf, format="PNG")
                vb64 = base64.standard_b64encode(vbuf.getvalue()).decode()

                verify_resp = client.messages.create(
                    model=CLAUDE_MODEL,
                    max_tokens=256,
                    temperature=0.0,
                    messages=[{"role": "user", "content": [
                        {"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": vb64}},
                        {"type": "text", "text": (
                            f"Look at {s['question']} on this exam page. "
                            f"What letter is circled or marked as the student's answer? "
                            f"If any answer is crossed out with a slash or X, IGNORE it — "
                            f"only report the cleanly circled answer. "
                            f"Reply with ONLY the letter (A, B, C, or D) or UNCLEAR if you cannot tell."
                        )},
                    ]}],
                )
                verify_text = ""
                for block in verify_resp.content:
                    if block.type == "text":
                        verify_text = block.text.strip().upper()
                        break
                verified_letter = verify_text if verify_text in ("A", "B", "C", "D", "UNCLEAR") else "UNCLEAR"

                if verified_letter == "UNCLEAR" or verified_letter != original_letter:
                    _log.info("[MC-VERIFY] %s %s: original read '%s', verification read '%s' - flagging",
                              exam_row["anon_id"], s["question"], original_letter, verified_letter)
                    mc_verified.append({
                        "question": s["question"],
                        "original_letter": original_letter,
                        "verified_letter": verified_letter,
                        "page": page,
                    })
                else:
                    _log.info("[MC-VERIFY] %s %s: confirmed '%s'",
                              exam_row["anon_id"], s["question"], original_letter)
            except Exception as e:
                _log.warning("[MC-VERIFY] %s %s: verification failed - %s",
                             exam_row["anon_id"], s["question"], e)

    # --- Collect review flags ---
    # Handwriting flags from the model + contradiction-resolved questions for professor review.
    flags = []
    if data.get("scores"):
        for s in data["scores"]:
            if s.pop("handwriting_flag", False):
                flags.append({"question": s["question"], "reason": "handwriting", "page": s.get("page")})
                _log.info("[HANDWRITING] %s %s flagged for handwriting review",
                          exam_row["anon_id"], s["question"])
            elif _HANDWRITING_UNCERTAINTY.search(s.get("feedback", "")):
                flags.append({"question": s["question"], "reason": "handwriting", "page": s.get("page")})
                _log.info("[HANDWRITING] %s %s flagged via feedback scan",
                          exam_row["anon_id"], s["question"])
    for q in contradiction_resolved:
        # Look up page from scores list for contradiction flags
        q_page = next((s.get("page") for s in data.get("scores", []) if s["question"] == q), None)
        flags.append({"question": q, "reason": "contradiction_resolved", "page": q_page})
    for mc in mc_verified:
        flags.append({
            "question": mc["question"],
            "reason": "mc_letter_mismatch",
            "page": mc["page"],
            "original_letter": mc["original_letter"],
            "verified_letter": mc["verified_letter"],
        })
    if flags:
        data["review_flags"] = flags

    # Phase 2: recalc + hard cap + round + letter grade (shared with audit)
    finalize_scores(data)

    return data


def _boundary_regrade(exam_row, rubric, original_data: dict) -> dict:
    """Re-grade an exam near a letter grade boundary and reconcile the two passes.

    If both passes agree on the letter grade, keep the original.
    If they disagree, average the earned scores to eliminate single-pass noise.
    The boundary_check field is added to grade_data for audit trail.
    """
    orig_earned = original_data["total_earned"]
    orig_possible = original_data["total_possible"]
    orig_pct = orig_earned / orig_possible * 100
    orig_grade = original_data["letter_grade"]

    _log.info("[BOUNDARY] %s scored %.1f%% (%s) - triggering verification re-grade",
              exam_row["anon_id"], orig_pct, orig_grade)

    regrade_data = _grade_exam(exam_row, rubric)
    regrade_earned = regrade_data["total_earned"]
    regrade_pct = regrade_earned / regrade_data["total_possible"] * 100
    regrade_grade = regrade_data["letter_grade"]

    if orig_grade == regrade_grade:
        _log.info("[BOUNDARY] %s confirmed: both passes agree on %s (%.1f%% vs %.1f%%)",
                  exam_row["anon_id"], orig_grade, orig_pct, regrade_pct)
        original_data["boundary_check"] = {
            "pass_1": {"earned": orig_earned, "pct": round(orig_pct, 1), "grade": orig_grade},
            "pass_2": {"earned": regrade_earned, "pct": round(regrade_pct, 1), "grade": regrade_grade},
            "result": "confirmed",
        }
        return original_data

    # Letter grades disagree - average the earned scores
    avg_earned = round((orig_earned + regrade_earned) / 2, 2)
    avg_pct = avg_earned / orig_possible * 100
    final_grade = letter_grade(avg_pct)

    _log.info("[BOUNDARY] %s disagreement: pass 1 = %s (%.1f%%), pass 2 = %s (%.1f%%) "
              "- averaged to %s (%.1f%%)",
              exam_row["anon_id"], orig_grade, orig_pct,
              regrade_grade, regrade_pct, final_grade, avg_pct)

    original_data["total_earned"] = avg_earned
    original_data["letter_grade"] = final_grade
    original_data["boundary_check"] = {
        "pass_1": {"earned": orig_earned, "pct": round(orig_pct, 1), "grade": orig_grade},
        "pass_2": {"earned": regrade_earned, "pct": round(regrade_pct, 1), "grade": regrade_grade},
        "result": "averaged",
    }
    return original_data


def _grade_one_worker(anon_id: str):
    """Grade a single exam. Called by ThreadPoolExecutor — each call gets its own DB connection."""
    if _grade_abort.is_set():
        return
    conn = sqlite3.connect(str(DB_PATH), detect_types=sqlite3.PARSE_DECLTYPES)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    try:
        exam = conn.execute(
            "SELECT anon_id, version, file_path, grade_data FROM exams WHERE anon_id=?", (anon_id,)
        ).fetchone()
        if not exam:
            raise ValueError(f"Exam {anon_id} not found")

        # Resume safety: skip if already graded (e.g. by another route mid-batch)
        if exam["grade_data"] is not None:
            _log.info("Skipping %s — already graded", anon_id)
            with _grade_lock:
                _grade_job["done"] += 1
            return

        rubric = conn.execute(
            "SELECT * FROM rubrics WHERE version=? ORDER BY id DESC LIMIT 1",
            (exam["version"],)
        ).fetchone()
        if not rubric:
            raise ValueError(f"No rubric for Version {exam['version']}")

        grade_data = _grade_exam(exam, rubric)

        # Boundary verification: re-grade exams near letter grade thresholds
        if grade_data.get("total_possible", 0) > 0:
            pct = grade_data["total_earned"] / grade_data["total_possible"] * 100
            if _is_boundary_score(pct):
                grade_data = _boundary_regrade(exam, rubric, grade_data)

        conn.execute(
            "UPDATE exams SET grade_data=?, graded_at=? WHERE anon_id=?",
            (json.dumps(grade_data), datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S"), anon_id)
        )
        conn.commit()
        _log.info("Graded %s — %.1f/%d", anon_id,
                  grade_data.get("total_earned", 0), grade_data.get("total_possible", 0))
        with _grade_lock:
            _grade_job["done"] += 1
    except Exception as e:
        _log.error("Grading failed for %s: %s", anon_id, e)
        with _grade_lock:
            _grade_job["failed"] += 1
            _grade_job["errors"].append(f"{anon_id}: {str(e)[:120]}")
    finally:
        conn.close()


def _run_grade_pool(anon_ids: list):
    """Background thread: grades all exams using a thread pool."""
    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=GRADE_WORKERS) as pool:
            pool.map(_grade_one_worker, anon_ids)
    finally:
        with _grade_lock:
            _grade_job["running"] = False
            if _grade_abort.is_set():
                _grade_job["aborted"] = True


def _enqueue_and_start(anon_ids: list):
    """Start parallel grading for the given exam IDs."""
    global _grade_job
    with _grade_lock:
        if _grade_job["running"]:
            # Already running — don't start a second pool
            return
        _grade_abort.clear()
        _grade_job = {"running": True, "total": len(anon_ids),
                      "done": 0, "failed": 0, "errors": []}
    threading.Thread(target=_run_grade_pool, args=(anon_ids,), daemon=True).start()


@app.route("/grade/<anon_id>", methods=["POST"])
def grade_one(anon_id):
    db = get_db()
    exam = db.execute("SELECT anon_id, version, file_path, grade_data FROM exams WHERE anon_id=?", (anon_id,)).fetchone()
    if not exam:
        flash("Exam not found.", "danger")
        return redirect(url_for("exams"))

    rubric = get_rubric(exam["version"])
    if not rubric:
        flash(f"No rubric found for Version {exam['version']}.", "danger")
        return redirect(url_for("exams"))
    try:
        grade_data = _grade_exam(exam, rubric)

        # Boundary verification: re-grade exams near letter grade thresholds
        if grade_data.get("total_possible", 0) > 0:
            pct = grade_data["total_earned"] / grade_data["total_possible"] * 100
            if _is_boundary_score(pct):
                grade_data = _boundary_regrade(exam, rubric, grade_data)

        db.execute(
            "UPDATE exams SET grade_data=?, graded_at=? WHERE anon_id=?",
            (json.dumps(grade_data), datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S"), anon_id)
        )
        db.commit()
        flash(f"Exam {anon_id} graded successfully.", "success")
    except Exception as e:
        _log.error("grade_one failed for %s: %s", anon_id, e, exc_info=True)
        flash(f"Grading failed for {anon_id}. See grading.log for details.", "danger")

    return redirect(url_for("results"))


@app.route("/grade-all", methods=["POST"])
def grade_all():
    db = get_db()
    version_filter = request.form.get("version", "").strip().upper() or None
    try:
        batch_filter = int(request.form.get("batch", "").strip())
    except (ValueError, TypeError):
        batch_filter = None

    query  = "SELECT anon_id FROM exams WHERE grade_data IS NULL"
    params = []
    if version_filter:
        query += " AND version=?"
        params.append(version_filter)
    if batch_filter:
        query += " AND batch=?"
        params.append(batch_filter)

    rows     = db.execute(query, params).fetchall()
    exam_ids = [r["anon_id"] for r in rows]

    if not exam_ids:
        flash("No ungraded exams found.", "info")
        return redirect(url_for("exams"))

    _enqueue_and_start(exam_ids)
    return redirect(url_for("grade_progress"))


@app.route("/grade-selected", methods=["POST"])
def grade_selected():
    anon_ids = request.form.getlist("anon_ids")
    if not anon_ids:
        return jsonify({"error": "No exams selected."}), 400

    db = get_db()
    placeholders = ",".join("?" * len(anon_ids))
    valid_ids = [
        r["anon_id"] for r in db.execute(
            f"SELECT anon_id FROM exams WHERE anon_id IN ({placeholders})", anon_ids
        ).fetchall()
    ]
    if not valid_ids:
        return jsonify({"error": "None of the selected exams were found."}), 400

    _enqueue_and_start(valid_ids)
    return jsonify({"ok": True, "total": len(valid_ids)})


@app.route("/regrade-selected", methods=["POST"])
def regrade_selected():
    anon_ids = request.form.getlist("anon_ids")
    if not anon_ids:
        return jsonify({"error": "No exams selected."}), 400

    db = get_db()
    placeholders = ",".join("?" * len(anon_ids))
    valid_ids = [
        r["anon_id"] for r in db.execute(
            f"SELECT anon_id FROM exams WHERE anon_id IN ({placeholders})", anon_ids
        ).fetchall()
    ]
    if not valid_ids:
        return jsonify({"error": "None of the selected exams were found."}), 400

    # Clear existing grades and review status so _grade_one_worker won't skip them
    db.execute(
        f"UPDATE exams SET grade_data=NULL, graded_at=NULL, reviewed=0 WHERE anon_id IN ({placeholders})",
        valid_ids,
    )
    db.commit()

    _enqueue_and_start(valid_ids)
    return jsonify({"ok": True, "total": len(valid_ids)})


@app.route("/grade-all/progress")
def grade_progress():
    return render_template("grade_progress.html")


@app.route("/grade-all/status")
def grade_status():
    with _grade_lock:
        return jsonify(dict(_grade_job))


@app.route("/grade-all/abort", methods=["POST"])
def grade_abort():
    _grade_abort.set()
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Routes: Exam detail / review
# ---------------------------------------------------------------------------

@app.route("/exam/<anon_id>")
def exam_detail(anon_id):
    db   = get_db()
    exam = db.execute("SELECT * FROM exams WHERE anon_id=?", (anon_id,)).fetchone()
    if not exam:
        flash("Exam not found.", "danger")
        return redirect(url_for("exams"))
    doc        = pdfium.PdfDocument(exam["file_path"])
    page_count = len(doc)
    doc.close()
    return render_template("exam_detail.html", exam=exam, page_count=page_count)


@app.route("/exam/<anon_id>/clear-grade", methods=["POST"])
def clear_grade(anon_id):
    db = get_db()
    db.execute(
        "UPDATE exams SET grade_data=NULL, graded_at=NULL WHERE anon_id=?", (anon_id,)
    )
    db.commit()
    flash(f"Grade cleared for {anon_id}.", "success")
    return _safe_redirect_back("exams")


@app.route("/clear-all-grades", methods=["POST"])
def clear_all_grades():
    version_filter = request.form.get("version", "").strip().upper() or None
    try:
        batch_filter = int(request.form.get("batch", "").strip())
    except (ValueError, TypeError):
        batch_filter = None
    db = get_db()
    query  = "UPDATE exams SET grade_data=NULL, graded_at=NULL WHERE grade_data IS NOT NULL"
    params = []
    if version_filter:
        query += " AND version=?"
        params.append(version_filter)
    if batch_filter:
        query += " AND batch=?"
        params.append(batch_filter)
    db.execute(query, params)
    count = db.execute("SELECT changes()").fetchone()[0]
    db.commit()
    flash(f"Cleared grades for {count} exam(s). Ready to re-grade.", "success")
    return redirect(url_for("exams"))


@app.route("/exam/<anon_id>/update-name", methods=["POST"])
def update_exam_name(anon_id):
    db   = get_db()
    name = request.form.get("student_name", "").strip()
    sid  = request.form.get("student_sid",  "").strip()
    if not name:
        flash("Name cannot be empty.", "danger")
        return redirect(url_for("exam_detail", anon_id=anon_id))
    db.execute(
        "UPDATE exams SET student_name=?, student_sid=? WHERE anon_id=?",
        (name, sid, anon_id)
    )
    db.commit()
    flash("Student info updated.", "success")
    return redirect(url_for("exam_detail", anon_id=anon_id))


@app.route("/exam/<anon_id>/page/<int:page_num>")
def exam_page_image(anon_id, page_num):
    db   = get_db()
    exam = db.execute("SELECT file_path FROM exams WHERE anon_id=?", (anon_id,)).fetchone()
    if not exam:
        return "Not found", 404
    try:
        img_bytes = _render_page_png(exam["file_path"], page_num)
        return send_file(io.BytesIO(img_bytes), mimetype="image/png")
    except Exception as e:
        _log.error("Error rendering page %d for %s: %s", page_num, anon_id, e)
        return "Error rendering page", 500


# ---------------------------------------------------------------------------
# Routes: Results
# ---------------------------------------------------------------------------

@app.route("/results")
def results():
    db    = get_db()
    show  = request.args.get("show_names") == "1"
    version_filter = request.args.get("version", "").strip().upper() or None
    try:
        batch_filter = int(request.args.get("batch", "").strip())
    except (ValueError, TypeError):
        batch_filter = None

    query  = "SELECT * FROM exams WHERE grade_data IS NOT NULL"
    params = []
    if version_filter:
        query += " AND version=?"
        params.append(version_filter)
    if batch_filter:
        query += " AND batch=?"
        params.append(batch_filter)
    query += " ORDER BY version, batch, anon_id"

    rows = db.execute(query, params).fetchall()

    graded = []
    for row in rows:
        gd = json.loads(row["grade_data"])
        graded.append({
            "anon_id":      row["anon_id"],
            "student_name": row["student_name"] if show else None,
            "student_sid":  row["student_sid"]  if show else None,
            "version":      row["version"],
            "batch":        row["batch"],
            "total_earned": gd.get("total_earned", "?"),
            "total_possible": gd.get("total_possible", "?"),
            "letter_grade": gd.get("letter_grade", "?"),
            "overall_feedback": gd.get("overall_feedback", ""),
            "scores":       gd.get("scores", []),
            "boundary_check": gd.get("boundary_check"),
            "review_flags": gd.get("review_flags", []),
            "reviewed":     bool(row["reviewed"]),
        })

    versions = db.execute(
        "SELECT DISTINCT version FROM exams WHERE grade_data IS NOT NULL ORDER BY version"
    ).fetchall()
    batches = db.execute(
        "SELECT DISTINCT batch FROM exams WHERE grade_data IS NOT NULL ORDER BY batch"
    ).fetchall()

    # Compute audit health for suggestions
    audit_health = None
    try:
        comparisons, meta = load_audit_data()
        if comparisons:
            agg = compute_aggregate(comparisons)
            failures = []
            if agg["exact_match_pct"] < 70:
                failures.append(f"Exact Score Match {agg['exact_match_pct']:.0f}% (threshold: 70%)")
            if agg["within_1_pct"] < 95:
                failures.append(f"Within-1-Point {agg['within_1_pct']:.0f}% (threshold: 95%)")
            if agg["mae"] >= 1.0:
                failures.append(f"MAE {agg['mae']:.2f} pts (threshold: <1.0)")
            if agg["grade_agreement_pct"] < 80:
                failures.append(f"Grade Agreement {agg['grade_agreement_pct']:.0f}% (threshold: 80%)")
            # Find questions with largest average deviation, keyed by (version, question)
            from collections import defaultdict
            q_diffs = defaultdict(list)
            for c in comparisons:
                ver = c.get("version", "?")
                for q in c["per_question"]:
                    if not q["question"].upper().startswith("TOTA"):
                        q_diffs[(ver, q["question"])].append(q["abs_diff"])
            worst_qs = []
            if q_diffs:
                ranked = sorted(q_diffs, key=lambda k: sum(q_diffs[k]) / len(q_diffs[k]), reverse=True)
                for ver, qname in ranked[:3]:
                    avg = sum(q_diffs[(ver, qname)]) / len(q_diffs[(ver, qname)])
                    worst_qs.append(f"{ver} {qname} (avg {avg:.1f} pt)")
            audit_health = {"failures": failures, "worst_questions": worst_qs}
    except Exception:
        pass

    return render_template("results.html",
                           graded=graded,
                           show_names=show,
                           versions=[r["version"] for r in versions],
                           batches=[r["batch"] for r in batches],
                           version_filter=version_filter,
                           batch_filter=batch_filter,
                           audit_health=audit_health)


@app.route("/exam/<anon_id>/toggle-reviewed", methods=["POST"])
def toggle_reviewed(anon_id):
    db = get_db()
    exam = db.execute("SELECT reviewed FROM exams WHERE anon_id=?", (anon_id,)).fetchone()
    if not exam:
        return jsonify({"error": "Exam not found"}), 404
    new_val = 0 if exam["reviewed"] else 1
    db.execute("UPDATE exams SET reviewed=? WHERE anon_id=?", (new_val, anon_id))
    db.commit()
    return jsonify({"reviewed": bool(new_val)})


# ---------------------------------------------------------------------------
def _csv_safe_id(anon_id: str) -> str:
    """Prevent Excel formula interpretation for IDs starting with - + = @."""
    if anon_id and anon_id[0] in "-+=@":
        return f'="{anon_id}"'
    return anon_id

# Routes: Export CSV
# ---------------------------------------------------------------------------

@app.route("/export")
def export():
    db = get_db()
    rows = db.execute(
        "SELECT * FROM exams WHERE grade_data IS NOT NULL ORDER BY version, batch, anon_id"
    ).fetchall()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "Anon ID", "Student Name", "Student SID", "Version", "Batch",
        "Total Earned", "Total Possible", "Percentage", "Letter Grade",
        "Overall Feedback"
    ])

    for row in rows:
        gd  = json.loads(row["grade_data"])
        earned   = gd.get("total_earned", 0)
        possible = gd.get("total_possible", 0)
        pct = f"{(earned/possible*100):.1f}%" if possible else "N/A"
        writer.writerow([
            _csv_safe_id(row["anon_id"]),
            row["student_name"],
            row["student_sid"] or "",
            row["version"],
            row["batch"],
            earned,
            possible,
            pct,
            gd.get("letter_grade", ""),
            gd.get("overall_feedback", ""),
        ])

    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode()),
        mimetype="text/csv",
        as_attachment=True,
        download_name=f"grades_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    )


# ---------------------------------------------------------------------------
# Routes: Batch PDF upload & split
# ---------------------------------------------------------------------------

@app.route("/upload-batch", methods=["GET", "POST"])
def upload_batch():
    if request.method == "POST":
        version        = request.form.get("version", "").strip().upper()
        pages_per_exam = request.form.get("pages_per_exam", "").strip()

        if not version:
            flash("Exam version is required.", "danger")
            return redirect(url_for("upload_batch"))

        try:
            pages_per_exam = int(pages_per_exam)
            if pages_per_exam < 1:
                raise ValueError
        except ValueError:
            flash("Pages per exam must be a whole number of 1 or more.", "danger")
            return redirect(url_for("upload_batch"))

        # Process up to 3 batch slots — each gets its own batch number
        all_exams = []
        for slot in range(1, 4):
            file = request.files.get(f"batch_pdf_{slot}")
            if not file or not file.filename:
                continue
            try:
                batch_num = int(request.form.get(f"batch_{slot}", slot))
            except ValueError:
                batch_num = slot

            pdf_bytes = file.read(MAX_UPLOAD_BYTES + 1)
            if len(pdf_bytes) > MAX_UPLOAD_BYTES:
                flash(f"Slot {slot} PDF exceeds 50 MB limit — skipped.", "warning")
                continue
            if not pdf_bytes:
                flash(f"Slot {slot} PDF is empty — skipped.", "warning")
                continue

            try:
                exams = split_by_page_count_bytes(pdf_bytes, pages_per_exam, UPLOAD_DIR)
            except Exception as e:
                flash(f"Failed to split slot {slot}: {e}", "danger")
                continue

            for exam in exams:
                exam["batch"] = batch_num
            all_exams.extend(exams)

        if not all_exams:
            flash("No valid PDFs were uploaded.", "danger")
            return redirect(url_for("upload_batch"))

        # Check cover page consistency across all exams (local, no API)
        check_cover_consistency(all_exams)

        # Discard any existing review sessions before creating a new one
        for stale in DATA_DIR.glob("review_*.json"):
            try:
                stale_data = json.loads(stale.read_text())
                for e in stale_data.get("exams", []):
                    Path(e["file_path"]).unlink(missing_ok=True)
            except Exception:
                pass
            stale.unlink(missing_ok=True)

        # Store split results immediately — OCR runs in background if opted in
        review_id   = secrets.token_hex(8)
        review_data = {"version": version, "exams": all_exams}
        review_path = DATA_DIR / f"review_{review_id}.json"
        review_path.write_text(json.dumps(review_data))

        if request.form.get("run_ocr") == "1":
            t = threading.Thread(
                target=_run_ocr_background,
                args=(review_id, review_path),
                daemon=True,
            )
            t.start()

        return redirect(url_for("batch_review", review_id=review_id))

    return render_template("upload_batch.html")


@app.route("/upload-batch/review/<review_id>")
def batch_review(review_id):
    if not _valid_review_id(review_id):
        flash("Invalid session.", "danger")
        return redirect(url_for("upload_batch"))
    review_path = DATA_DIR / f"review_{review_id}.json"
    if not review_path.exists():
        flash("Review session not found. Please upload again.", "danger")
        return redirect(url_for("upload_batch"))
    data         = json.loads(review_path.read_text())
    db           = get_db()
    roster_count = db.execute("SELECT COUNT(*) FROM roster").fetchone()[0]
    with _ocr_jobs_lock:
        ocr_running = _ocr_jobs.get(review_id, {}).get("running", False)
    return render_template("batch_review.html",
                           review_id=review_id,
                           version=data["version"],
                           exams=data["exams"],
                           roster_count=roster_count,
                           ocr_running=ocr_running)


@app.route("/upload-batch/ocr-status/<review_id>")
def ocr_status(review_id):
    if not _valid_review_id(review_id):
        return jsonify({"error": "Invalid session."}), 400
    with _ocr_jobs_lock:
        job = dict(_ocr_jobs.get(review_id, {"total": 0, "done": 0, "running": False}))
    review_path = DATA_DIR / f"review_{review_id}.json"
    exams = []
    if review_path.exists():
        try:
            data  = json.loads(review_path.read_text())
            exams = [{"name": e.get("name", ""), "sid": e.get("sid", ""),
                      "roster_score": e.get("roster_score", 0)}
                     for e in data["exams"]]
        except Exception:
            pass
    return jsonify({**job, "exams": exams})


@app.route("/upload-batch/ocr-abort/<review_id>", methods=["POST"])
def ocr_abort(review_id):
    if not _valid_review_id(review_id):
        return jsonify({"error": "Invalid session."}), 400
    with _ocr_jobs_lock:
        if review_id in _ocr_jobs and _ocr_jobs[review_id].get("running"):
            _ocr_jobs[review_id]["aborted"] = True
            return jsonify({"ok": True, "message": "Abort signal sent — stopping after current exam."})
    return jsonify({"ok": False, "message": "No active OCR job for this session."})


@app.route("/upload-batch/confirm/<review_id>", methods=["POST"])
def batch_confirm(review_id):
    if not _valid_review_id(review_id):
        flash("Invalid session.", "danger")
        return redirect(url_for("upload_batch"))
    review_path = DATA_DIR / f"review_{review_id}.json"
    if not review_path.exists():
        flash("Review session not found. Please upload again.", "danger")
        return redirect(url_for("upload_batch"))

    data  = json.loads(review_path.read_text())
    db    = get_db()
    saved = 0

    for i, exam in enumerate(data["exams"]):
        student_name = request.form.get(f"name_{i}", "").strip()
        student_sid  = request.form.get(f"sid_{i}",  "").strip()
        anon_id      = generate_anon_id()
        old_path     = Path(exam["file_path"])
        new_path     = UPLOAD_DIR / f"{anon_id}.pdf"
        if old_path.exists():
            # Copy then delete — avoids WinError 32 if a thumbnail request
            # still has the file open in pypdfium2
            new_path.write_bytes(old_path.read_bytes())
            old_path.unlink(missing_ok=True)
        db.execute(
            "INSERT INTO exams (anon_id, student_name, student_sid, version, batch, file_path) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (anon_id, student_name, student_sid, data["version"], exam["batch"], str(new_path))
        )
        saved += 1

    db.commit()
    review_path.unlink(missing_ok=True)
    flash(
        f"{saved} exams saved → Version {data['version']}.",
        "success"
    )
    return redirect(url_for("exams"))


@app.route("/upload-batch/discard/<review_id>", methods=["POST"])
def batch_discard(review_id):
    if not _valid_review_id(review_id):
        flash("Invalid session.", "danger")
        return redirect(url_for("upload_batch"))
    review_path = DATA_DIR / f"review_{review_id}.json"
    if review_path.exists():
        try:
            data = json.loads(review_path.read_text())
            for exam in data.get("exams", []):
                Path(exam["file_path"]).unlink(missing_ok=True)
        except Exception:
            pass
        review_path.unlink(missing_ok=True)
    with _ocr_jobs_lock:
        _ocr_jobs.pop(review_id, None)
    return redirect(url_for("upload_batch"))


@app.route("/upload-batch/preview/<review_id>/<int:exam_index>/<int:page_num>")
def batch_preview_page(review_id, exam_index, page_num):
    """Serve a rendered page image from a temp split exam (before it's saved to DB)."""
    if not _valid_review_id(review_id):
        return "Invalid session", 400
    review_path = DATA_DIR / f"review_{review_id}.json"
    if not review_path.exists():
        return "Session expired", 404
    data = json.loads(review_path.read_text())
    if exam_index < 0 or exam_index >= len(data["exams"]):
        return "Not found", 404
    file_path = data["exams"][exam_index]["file_path"]
    # Use lower scale for thumbnails to keep them fast
    scale = 0.75 if request.args.get("thumb") else 1.5
    try:
        img_bytes = _render_page_png(file_path, page_num, scale=scale)
        return send_file(io.BytesIO(img_bytes), mimetype="image/png")
    except Exception as e:
        _log.error("Error rendering preview page %d for review %s exam %d: %s", page_num, review_id, exam_index, e)
        return "Error rendering page", 500


# ---------------------------------------------------------------------------
# Routes: Roster management (local only — never sent to Claude)
# ---------------------------------------------------------------------------

@app.route("/roster", methods=["GET", "POST"])
def roster():
    db = get_db()
    if request.method == "POST":
        file = request.files.get("roster_csv")
        if not file or not file.filename:
            flash("Please select a CSV file to upload.", "danger")
            return redirect(url_for("roster"))

        try:
            content  = file.read().decode("utf-8-sig")  # strip BOM if present
            reader   = csv.reader(io.StringIO(content))
            rows     = list(reader)
        except Exception as e:
            flash(f"Could not read CSV: {e}", "danger")
            return redirect(url_for("roster"))

        # Auto-detect and skip header row
        entries = []
        for i, row in enumerate(rows):
            if len(row) < 2:
                continue
            first = row[0].strip()
            last  = row[1].strip()
            sid   = row[2].strip() if len(row) >= 3 else ""
            email = row[3].strip() if len(row) >= 4 else ""
            # Skip header row (first row whose first cell looks like a label)
            if i == 0 and first.lower() in ("first", "first_name", "firstname", "first name"):
                continue
            if not first and not last:
                continue
            entries.append((first, last, sid, email))

        if not entries:
            flash("No valid rows found in CSV (need at least first_name, last_name columns).", "danger")
            return redirect(url_for("roster"))

        db.execute("DELETE FROM roster")
        db.executemany("INSERT INTO roster (first_name, last_name, sid, email) VALUES (?, ?, ?, ?)", entries)
        db.commit()
        flash(f"Loaded {len(entries)} students from roster.", "success")
        return redirect(url_for("roster"))

    count      = db.execute("SELECT COUNT(*) FROM roster").fetchone()[0]
    preview    = db.execute("SELECT * FROM roster LIMIT 5").fetchall()
    show_names = request.args.get("show_names") == "1"
    return render_template("roster.html", count=count, preview=preview, show_names=show_names)


@app.route("/roster/clear", methods=["POST"])
def roster_clear():
    db = get_db()
    db.execute("DELETE FROM roster")
    db.commit()
    flash("Roster cleared.", "success")
    return redirect(url_for("roster"))


@app.route("/upload-batch/review/<review_id>/match-roster", methods=["POST"])
def match_roster(review_id):
    if not _valid_review_id(review_id):
        return jsonify({"error": "Invalid session."}), 400
    review_path = DATA_DIR / f"review_{review_id}.json"
    if not review_path.exists():
        return jsonify({"error": "Review session not found"}), 404

    data    = json.loads(review_path.read_text())
    db      = get_db()
    rows    = db.execute("SELECT * FROM roster").fetchall()
    entries = [dict(r) for r in rows]

    results = []
    for i, exam in enumerate(data["exams"]):
        best, score = match_against_roster(exam.get("name", ""), exam.get("sid", ""), entries)
        if best:
            tier = "high" if score >= 0.85 else ("medium" if score >= 0.60 else "low")
            results.append({
                "exam_index": i,
                "name":  f"{best['first_name']} {best['last_name']}",
                "sid":   best["sid"] or "",
                "score": score,
                "tier":  tier,
            })
        else:
            results.append({
                "exam_index": i,
                "name":  exam.get("name", ""),
                "sid":   exam.get("sid",  ""),
                "score": 0.0,
                "tier":  "low",
            })

    # Flag duplicates: multiple exams matched to the same roster student
    seen = {}  # (name, sid) -> [exam indices]
    for r in results:
        if r["score"] >= 0.60:
            key = (r["name"].strip().lower(), r["sid"].strip().lower())
            seen.setdefault(key, []).append(r["exam_index"])
    dup_map = {}
    for indices in seen.values():
        if len(indices) > 1:
            for idx in indices:
                dup_map[str(idx)] = indices

    return jsonify({"results": results, "duplicates": dup_map})


# ---------------------------------------------------------------------------
# Routes: Documentation
# ---------------------------------------------------------------------------

@app.route("/docs")
def docs():
    db = get_db()

    total_exams  = db.execute("SELECT COUNT(*) FROM exams").fetchone()[0]
    graded_exams = db.execute("SELECT COUNT(*) FROM exams WHERE grade_data IS NOT NULL").fetchone()[0]
    exam_versions = [r[0] for r in db.execute(
        "SELECT DISTINCT version FROM exams ORDER BY version"
    ).fetchall()]
    rubric_versions = [r[0] for r in db.execute(
        "SELECT DISTINCT version FROM rubrics ORDER BY version"
    ).fetchall()]
    batches = [r[0] for r in db.execute(
        "SELECT DISTINCT batch FROM exams ORDER BY batch"
    ).fetchall()]

    # Estimate cost: ~$0.12 average per graded exam (midpoint of $0.09–$0.15, 12-page exam)
    est_cost = round(graded_exams * 0.12, 2)

    # Rubric files on disk
    rubric_files = list(RUBRIC_DIR.glob("*"))
    rubric_file_count = len(rubric_files)

    # Upload dir size (MB)
    upload_bytes = sum(f.stat().st_size for f in UPLOAD_DIR.glob("*.pdf") if f.is_file())
    upload_mb    = round(upload_bytes / 1_048_576, 1)

    # Build info written by pre-commit hook
    build_info_path = BASE_DIR / "build_info.json"
    build_info = {}
    if build_info_path.exists():
        try:
            build_info = json.loads(build_info_path.read_text())
        except Exception:
            pass

    return render_template("docs.html",
        total_exams=total_exams,
        graded_exams=graded_exams,
        ungraded_exams=total_exams - graded_exams,
        exam_versions=exam_versions,
        rubric_versions=rubric_versions,
        batches=batches,
        est_cost=est_cost,
        rubric_file_count=rubric_file_count,
        upload_mb=upload_mb,
        claude_model=CLAUDE_MODEL,
        ollama_model=OLLAMA_VISION_MODEL,
        build_info=build_info,
    )


# ---------------------------------------------------------------------------
# Routes: Per-student printable report
# ---------------------------------------------------------------------------

@app.route("/exam/<anon_id>/report/update", methods=["POST"])
def update_report(anon_id):
    db   = get_db()
    exam = db.execute("SELECT anon_id, grade_data FROM exams WHERE anon_id=?", (anon_id,)).fetchone()
    if not exam or not exam["grade_data"]:
        return jsonify({"error": "No graded data for this exam"}), 404

    gd = json.loads(exam["grade_data"])

    edited_questions = set()
    for i, s in enumerate(gd.get("scores", [])):
        raw = request.form.get(f"earned_{i}", "")
        try:
            val = float(raw)
            old_val = float(s["earned_points"])
            # Clamp to [0, max_points]
            s["earned_points"] = max(0.0, min(val, float(s["max_points"])))
            if abs(val - old_val) > 0.001:
                edited_questions.add(s["question"])
        except (ValueError, TypeError):
            pass  # leave original if the field was blank or invalid

        # Update feedback if provided
        fb = request.form.get(f"feedback_{i}")
        if fb is not None:
            s["feedback"] = fb.strip()

    # Clear review_flags for questions whose scores were manually edited
    if edited_questions and gd.get("review_flags"):
        gd["review_flags"] = [f for f in gd["review_flags"]
                              if f["question"] not in edited_questions]
        if not gd["review_flags"]:
            del gd["review_flags"]

    # Recalculate totals
    total_earned   = sum(float(s["earned_points"]) for s in gd["scores"])
    total_possible = float(gd["total_possible"])
    total_earned   = min(round(total_earned, 2), total_possible)  # hard cap
    pct            = total_earned / total_possible * 100 if total_possible else 0

    gd["total_earned"]  = total_earned
    gd["letter_grade"]  = letter_grade(pct)

    db.execute("UPDATE exams SET grade_data=? WHERE anon_id=?",
               (json.dumps(gd), anon_id))
    db.commit()
    return jsonify({
        "total_earned":   total_earned,
        "total_possible": gd["total_possible"],
        "pct":            round(pct, 1),
        "letter_grade":   gd["letter_grade"]
    })


@app.route("/exam/<anon_id>/dismiss-flag", methods=["POST"])
def dismiss_flag(anon_id):
    db = get_db()
    exam = db.execute("SELECT anon_id, grade_data FROM exams WHERE anon_id=?", (anon_id,)).fetchone()
    if not exam or not exam["grade_data"]:
        return jsonify({"error": "Exam not found"}), 404

    gd = json.loads(exam["grade_data"])
    body = request.get_json(silent=True) or {}
    question = body.get("question", "")

    flags = gd.get("review_flags", [])
    gd["review_flags"] = [f for f in flags if f["question"] != question]
    remaining = len(gd["review_flags"])
    if not gd["review_flags"]:
        del gd["review_flags"]

    db.execute("UPDATE exams SET grade_data=? WHERE anon_id=?",
               (json.dumps(gd), anon_id))
    db.commit()
    return jsonify({"ok": True, "remaining_flags": remaining})


@app.route("/exam/<anon_id>/report")
def student_report(anon_id):
    db   = get_db()
    exam = db.execute("SELECT * FROM exams WHERE anon_id=?", (anon_id,)).fetchone()
    if not exam or not exam["grade_data"]:
        flash("No graded data for this exam.", "warning")
        return redirect(url_for("results"))
    gd        = json.loads(exam["grade_data"])
    pct       = gd["total_earned"] / gd["total_possible"] * 100 if gd["total_possible"] else 0
    show_name    = request.args.get("show_name") == "1"
    include_scans = request.args.get("scans") == "1"
    page_count   = 0
    if include_scans:
        doc = pdfium.PdfDocument(exam["file_path"])
        page_count = len(doc)
        doc.close()
    return render_template("student_report.html", exam=exam, gd=gd, pct=pct,
                           show_name=show_name, page_count=page_count)


# ---------------------------------------------------------------------------
# Routes: Detailed CSV export (per-question columns)
# ---------------------------------------------------------------------------

@app.route("/export-detailed")
def export_detailed():
    db = get_db()
    version_filter = request.args.get("version", "").upper() or None
    try:
        batch_filter = int(request.args.get("batch", ""))
    except (ValueError, TypeError):
        batch_filter = None

    query  = "SELECT * FROM exams WHERE grade_data IS NOT NULL"
    params = []
    if version_filter:
        query += " AND version=?"
        params.append(version_filter)
    if batch_filter:
        query += " AND batch=?"
        params.append(batch_filter)
    query += " ORDER BY version, batch, anon_id"

    rows = db.execute(query, params).fetchall()

    # Collect all unique question names in encounter order
    all_questions: dict = {}
    parsed = []
    for row in rows:
        gd = json.loads(row["grade_data"])
        parsed.append((row, gd))
        for s in gd.get("scores", []):
            q = s.get("question", "")
            if q and q not in all_questions:
                all_questions[q] = None

    question_names = list(all_questions.keys())

    output = io.StringIO()
    writer = csv.writer(output)

    # Header row
    fixed_cols = [
        "Anon ID", "Name", "SID", "Version", "Batch",
        "Total Earned", "Total Possible", "Pct", "Grade", "Overall Feedback"
    ]
    dynamic_cols = []
    for q in question_names:
        dynamic_cols += [f"{q} Earned", f"{q} Max", f"{q} Feedback"]
    writer.writerow(fixed_cols + dynamic_cols)

    # Data rows
    for row, gd in parsed:
        earned   = gd.get("total_earned", 0)
        possible = gd.get("total_possible", 0)
        pct      = f"{(earned / possible * 100):.1f}%" if possible else "N/A"

        # Build lookup by question name
        scores_by_q = {s["question"]: s for s in gd.get("scores", [])}

        fixed = [
            _csv_safe_id(row["anon_id"]),
            row["student_name"],
            row["student_sid"] or "",
            row["version"],
            row["batch"],
            earned,
            possible,
            pct,
            gd.get("letter_grade", ""),
            gd.get("overall_feedback", ""),
        ]
        dynamic = []
        for q in question_names:
            s = scores_by_q.get(q, {})
            dynamic += [
                s.get("earned_points", ""),
                s.get("max_points", ""),
                s.get("feedback", ""),
            ]
        writer.writerow(fixed + dynamic)

    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode()),
        mimetype="text/csv",
        as_attachment=True,
        download_name=f"grades_detailed_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    )


# ---------------------------------------------------------------------------
# Routes: Analytics
# ---------------------------------------------------------------------------

@app.route("/analytics")
def analytics():
    db = get_db()
    version_filter = request.args.get("version", "").upper() or None
    try:
        batch_filter = int(request.args.get("batch", ""))
    except (ValueError, TypeError):
        batch_filter = None
    q_version      = request.args.get("q_version", "").upper() or None

    # Main query — drives distribution, band, and summary stats
    query  = "SELECT grade_data FROM exams WHERE grade_data IS NOT NULL"
    params = []
    if version_filter:
        query += " AND version=?"
        params.append(version_filter)
    if batch_filter:
        query += " AND batch=?"
        params.append(batch_filter)

    rows = db.execute(query, params).fetchall()

    pct_scores = []
    for row in rows:
        gd       = json.loads(row["grade_data"])
        possible = float(gd.get("total_possible", 0))
        earned   = float(gd.get("total_earned",   0))
        if possible > 0:
            pct_scores.append(round(earned / possible * 100, 2))

    # Question stats — separate query filtered by q_version only
    q_query  = "SELECT grade_data FROM exams WHERE grade_data IS NOT NULL"
    q_params = []
    if q_version:
        q_query += " AND version=?"
        q_params.append(q_version)
    elif version_filter:
        q_query += " AND version=?"
        q_params.append(version_filter)

    q_rows = db.execute(q_query, q_params).fetchall()
    question_stats = {}
    for row in q_rows:
        gd = json.loads(row["grade_data"])
        for s in gd.get("scores", []):
            q  = s.get("question", "").strip()
            ep = float(s.get("earned_points", 0))
            mp = float(s.get("max_points",   0))
            if not q:
                continue
            if q not in question_stats:
                question_stats[q] = {"earned": 0.0, "possible": 0.0, "n": 0, "below70": 0}
            question_stats[q]["earned"]   += ep
            question_stats[q]["possible"] += mp
            question_stats[q]["n"]        += 1
            if mp > 0 and ep / mp < 0.70:
                question_stats[q]["below70"] += 1

    # Grade distribution — 10 bins: [0,10), [10,20), …, [90,100]
    dist_bins   = [0] * 10
    dist_labels = ["0–10","10–20","20–30","30–40","40–50",
                   "50–60","60–70","70–80","80–90","90–100"]
    for p in pct_scores:
        dist_bins[min(int(p // 10), 9)] += 1

    # Letter-grade band counts
    grade_bands = {"A": 0, "B": 0, "C": 0, "D": 0, "F": 0}
    for p in pct_scores:
        grade_bands[letter_grade(p)] += 1

    # Per-question averages and fail rates
    q_labels   = list(question_stats.keys())
    q_avg_pct  = []
    q_fail_rate = []
    struggling  = []   # questions with avg < 60%
    for q in q_labels:
        st  = question_stats[q]
        avg = round(st["earned"] / st["possible"] * 100, 1) if st["possible"] else 0.0
        fr  = round(st["below70"] / st["n"] * 100, 1)       if st["n"]        else 0.0
        q_avg_pct.append(avg)
        q_fail_rate.append(fr)
        if avg < 60:
            struggling.append({"question": q, "avg": avg, "fail_rate": fr})

    # Summary statistics
    n = len(pct_scores)
    if n > 0:
        mean   = round(_stats.mean(pct_scores),   1)
        median = round(_stats.median(pct_scores), 1)
        stdev  = round(_stats.stdev(pct_scores),  1) if n > 1 else 0.0
        lo     = round(min(pct_scores), 1)
        hi     = round(max(pct_scores), 1)
        pass_r = round(sum(1 for p in pct_scores if p >= 70) / n * 100, 1)
    else:
        mean = median = stdev = lo = hi = pass_r = 0.0

    versions = db.execute(
        "SELECT DISTINCT version FROM exams WHERE grade_data IS NOT NULL ORDER BY version"
    ).fetchall()
    batches = db.execute(
        "SELECT DISTINCT batch FROM exams WHERE grade_data IS NOT NULL ORDER BY batch"
    ).fetchall()

    return render_template("analytics.html",
        n=n, mean=mean, median=median, stdev=stdev, lo=lo, hi=hi, pass_r=pass_r,
        dist_bins=dist_bins, dist_labels=dist_labels,
        grade_bands=grade_bands,
        q_labels=q_labels, q_avg_pct=q_avg_pct, q_fail_rate=q_fail_rate,
        struggling=struggling,
        versions=[r["version"] for r in versions],
        batches=[r["batch"] for r in batches],
        version_filter=version_filter,
        batch_filter=batch_filter,
        q_version=q_version,
        now=datetime.now(timezone.utc).strftime("%B %d, %Y"),
    )


# ---------------------------------------------------------------------------
# Launch Ollama
# ---------------------------------------------------------------------------

@app.route("/launch-ollama", methods=["POST"])
def launch_ollama():
    import subprocess, shutil, urllib.request
    ollama_path = shutil.which("ollama") or os.environ.get("OLLAMA_EXE")
    if not ollama_path:
        return {"ok": False, "message": "ollama not found in PATH. Set OLLAMA_EXE in .env."}, 500

    # Check if already running
    try:
        urllib.request.urlopen("http://localhost:11434/api/version", timeout=2)
        return {"ok": True, "message": "Ollama is already running."}
    except Exception:
        pass

    try:
        subprocess.Popen(
            [ollama_path, "serve"],
            creationflags=subprocess.CREATE_NEW_PROCESS_GROUP | subprocess.DETACHED_PROCESS,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        return {"ok": True, "message": "Ollama launched — allow a few seconds to load."}
    except Exception as e:
        return {"ok": False, "message": str(e)}, 500


# ---------------------------------------------------------------------------
# Routes: Audit
# ---------------------------------------------------------------------------

@app.route("/audit")
def audit_redirect():
    """Redirect legacy /audit bookmarks to /quality."""
    return redirect(url_for("quality_dashboard"), code=301)


@app.route("/quality")
def quality_dashboard():
    section = request.args.get("section", "audit")

    # Insights section — load insights.json and render
    if section == "insights":
        insights_path = AUDIT_DIR / "insights.json"
        insights_data = None
        if insights_path.exists():
            try:
                with open(insights_path) as f:
                    insights_data = json.load(f)
            except (json.JSONDecodeError, OSError):
                pass
        return render_template("quality.html", section="insights",
                               insights=insights_data)

    # Audit section (default) — requires proprietary audit modules
    if not _HAS_AUDIT:
        flash("Audit module not available.", "warning")
        return redirect(url_for("index"))

    db = get_db()
    graded_count = db.execute("SELECT COUNT(*) FROM exams WHERE grade_data IS NOT NULL").fetchone()[0]
    total_count  = db.execute("SELECT COUNT(*) FROM exams").fetchone()[0]

    # Current audit job status
    status = get_audit_status()

    # Available versions for filter dropdown
    versions = [r["version"] for r in db.execute("SELECT DISTINCT version FROM rubrics").fetchall()]

    # Build run history from audit JSON files — this is the single source of truth
    # for everything on this page (metrics, cumulative stats, progress bar)
    run_history = []
    active_files = []
    total_audited = 0
    if AUDIT_DIR.exists():
        for f in sorted(AUDIT_DIR.glob("audit_*.json"), reverse=True):
            try:
                with open(f) as fh:
                    data = json.load(fh)
                grade_matches = sum(1 for c in data["comparisons"] if c.get("grade_match"))
                n = len(data["comparisons"])
                run_history.append({
                    "filename": f.name,
                    "timestamp": data.get("timestamp", ""),
                    "model": data.get("audit_model", "?"),
                    "sample_size": n,
                    "agreement_pct": round(grade_matches / n * 100) if n else 0,
                })
                active_files.append(str(f))
                total_audited += n
            except (json.JSONDecodeError, KeyError):
                continue

    # Cumulative stats derived from the same file list
    cum_stats = {"total_audited": total_audited, "runs": len(active_files)}

    # Recommended target
    target_n = max(30, min(50, int(graded_count * 0.10)))
    remaining = max(0, target_n - total_audited)

    # Aggregate metrics from the same file list (no separate glob)
    metrics = None
    if active_files:
        comparisons, meta = load_audit_data(specific_files=active_files)
        if comparisons:
            metrics = compute_aggregate(comparisons)

    can_audit = graded_count >= 20 and not status["running"] and not _grade_job["running"]

    return render_template("quality.html",
                           section="audit",
                           graded_count=graded_count,
                           total_count=total_count,
                           cum_stats=cum_stats,
                           target_n=target_n,
                           remaining=remaining,
                           metrics=metrics,
                           status=status,
                           versions=versions,
                           run_history=run_history,
                           can_audit=can_audit)


def _require_audit():
    """Return an error response if audit modules are unavailable, else None."""
    if not _HAS_AUDIT:
        return jsonify({"error": "Audit module not available"}), 404
    return None


@app.route("/audit/confirm")
def audit_confirm():
    err = _require_audit()
    if err: return err
    sample_size = request.args.get("sample_size", 10, type=int)
    sample_size = min(max(sample_size, 1), 30)
    model_choice = request.args.get("model", "gemini")
    version = request.args.get("version", "").strip() or None

    if model_choice == "gemini":
        model_display = "Gemini 3.1 Pro (cross-family)"
        cost_low = sample_size * 0.05
        cost_high = sample_size * 0.15
        # Batches of 5 with 30s cooldown
        batches = (sample_size + 4) // 5
        minutes = max(2, batches * 1 + (batches - 1) * 0.5)
        time_estimate = f"~{minutes:.0f}-{minutes * 2:.0f} minutes"
    else:
        model_display = "Claude Opus 4.6 (within-family)"
        cost_low = sample_size * 0.20
        cost_high = sample_size * 0.50
        minutes = sample_size * 1.5
        time_estimate = f"~{minutes:.0f}-{minutes * 2:.0f} minutes"

    cost_estimate = f"~${cost_low:.2f}-${cost_high:.2f}"

    return render_template("audit_confirm.html",
                           sample_size=sample_size,
                           model_choice=model_choice,
                           model_display=model_display,
                           version=version,
                           cost_estimate=cost_estimate,
                           time_estimate=time_estimate)


@app.route("/audit/start", methods=["POST"])
def audit_start():
    err = _require_audit()
    if err: return err
    sample_size = request.form.get("sample_size", 10, type=int)
    version = request.form.get("version", "").strip() or None
    model_choice = request.form.get("model", "gemini")

    ok, msg = start_audit(sample_size=sample_size, version=version,
                          model_choice=model_choice)
    if not ok:
        flash(msg, "danger")
        return redirect(url_for("quality_dashboard"))

    flash(msg, "success")
    return redirect(url_for("quality_dashboard"))


@app.route("/audit/status")
def audit_status():
    err = _require_audit()
    if err: return err
    return jsonify(get_audit_status())


@app.route("/audit/abort", methods=["POST"])
def audit_abort():
    err = _require_audit()
    if err: return err
    _audit_abort.set()
    return jsonify({"ok": True})


@app.route("/audit/report", methods=["POST"])
def audit_report():
    err = _require_audit()
    if err: return err
    selected = request.form.getlist("selected")
    if not selected:
        flash("No audit runs selected. Check at least one run.", "warning")
        return redirect(url_for("quality_dashboard"))

    # Resolve to absolute paths and validate filenames
    specific_files = []
    for fname in selected:
        fp = AUDIT_DIR / fname
        if fp.exists() and fp.name.startswith("audit_") and fp.suffix == ".json":
            specific_files.append(str(fp))
    if not specific_files:
        flash("Selected audit files not found.", "danger")
        return redirect(url_for("quality_dashboard"))

    comparisons, meta = load_audit_data(specific_files=specific_files)
    if not comparisons:
        flash("No comparisons found in selected files.", "warning")
        return redirect(url_for("quality_dashboard"))

    report_path = DATA_DIR / "audit_report.pdf"
    build_report(comparisons, str(report_path), meta)
    return send_file(str(report_path), as_attachment=True,
                     download_name="rubrica_audit_report.pdf")


@app.route("/audit/archive", methods=["POST"])
def audit_archive():
    err = _require_audit()
    if err: return err
    selected = request.form.getlist("selected")
    if not selected:
        flash("No audit runs selected.", "warning")
        return redirect(url_for("quality_dashboard"))

    archive_dir = AUDIT_DIR / "archive"
    archive_dir.mkdir(exist_ok=True)
    moved = 0
    for fname in selected:
        fp = AUDIT_DIR / fname
        if fp.exists() and fp.name.startswith("audit_") and fp.suffix == ".json":
            fp.rename(archive_dir / fp.name)
            moved += 1

    flash(f"Archived {moved} audit run{'s' if moved != 1 else ''}.", "success")
    return redirect(url_for("quality_dashboard"))


@app.route("/quality/insights/refresh", methods=["POST"])
def insights_refresh():
    """Regenerate insights.json from all audit runs."""
    from insights import (load_all_audit_runs, deduplicate_comparisons,
                          compute_aggregate as insights_aggregate,
                          compute_per_question_cumulative, compute_trends,
                          build_run_history, load_rubric_metadata,
                          extract_structural_patterns, extract_feedback_patterns,
                          detect_bias_patterns, generate_lessons,
                          build_insights_json, save_insights)
    since = request.form.get("since", "").strip() or None
    try:
        runs = load_all_audit_runs(since=since)
        if not runs:
            flash("No audit runs found." + (f" (since {since})" if since else " Run an audit first."), "warning")
            return redirect(url_for("quality_dashboard", section="insights"))

        comparisons = deduplicate_comparisons(runs)
        cumulative = insights_aggregate(comparisons)
        per_question = compute_per_question_cumulative(comparisons)
        trends = compute_trends(runs)
        run_history = build_run_history(runs)
        rubric_meta = load_rubric_metadata()
        structural = extract_structural_patterns(per_question, rubric_meta)
        feedback = extract_feedback_patterns(comparisons)
        biases = detect_bias_patterns(comparisons)
        lessons = generate_lessons(structural, feedback, biases)

        data = build_insights_json(runs, comparisons, cumulative, per_question,
                                   trends, structural, feedback, biases,
                                   lessons, run_history)
        if since:
            data["since_filter"] = since
        save_insights(data)
        flash("Insights refreshed." + (f" (runs since {since})" if since else ""), "success")
    except Exception as e:
        _log.exception("Insights refresh failed")
        flash(f"Insights refresh failed: {e}", "danger")

    return redirect(url_for("quality_dashboard", section="insights"))


# ---------------------------------------------------------------------------
# Rubric Builder routes
# ---------------------------------------------------------------------------


def _save_enhanced_to_db(sid: str) -> tuple:
    """Build enhanced rubric from session, save to DB, clean up. Returns (version, other_versions) or raises."""
    enhanced = build_enhanced_rubric(sid)
    if not enhanced:
        return None, None

    version = enhanced["version"]
    db = get_db()

    db.execute(
        "UPDATE rubrics SET enhanced_rubric=? WHERE version=? AND id = "
        "(SELECT MAX(id) FROM rubrics WHERE version=?)",
        (json.dumps(enhanced, indent=2), version, version)
    )
    db.commit()
    delete_session(sid)
    discard_draft()

    other_versions = [
        r["version"] for r in db.execute(
            "SELECT DISTINCT version FROM rubrics WHERE version != ?", (version,)
        ).fetchall()
    ]
    return version, other_versions


@app.route("/rubric-builder/resume", methods=["POST"])
def builder_resume():
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    sess = load_draft()
    if not sess:
        return jsonify({"error": "No draft found"}), 404
    return jsonify({
        "session_id": sess["id"],
        "questions": sess["questions"],
        "total_q": len(sess["questions"]),
        "current_q": sess["current_q"],
        "phase": sess["phase"],
        "version": sess["version"],
    })


@app.route("/rubric-builder/discard", methods=["POST"])
def builder_discard():
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    info = get_draft_info()
    if info:
        delete_session(info["session_id"])
    discard_draft()
    return jsonify({"success": True})


@app.route("/rubric-builder/extract", methods=["POST"])
def builder_extract():
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    version = request.form.get("version", "").strip().upper()
    if not version:
        return jsonify({"error": "No version selected"}), 400
    rubric = get_rubric(version)
    if not rubric:
        return jsonify({"error": f"No rubric found for version {version}"}), 404

    sid = _new_session(version, rubric["content"], rubric["total_points"])
    sess = extract_questions(sid)

    if sess.get("error"):
        delete_session(sid)
        return jsonify({"error": sess["error"]}), 500

    return jsonify({
        "session_id": sid,
        "questions": sess["questions"],
        "total_q": len(sess["questions"]),
    })


@app.route("/rubric-builder/auto-enhance", methods=["POST"])
def builder_auto_enhance():
    """Run AI auto-enhancement, save to DB, return version for view/edit."""
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    body = request.get_json(silent=True) or {}
    sid = body.get("session_id")
    if not sid:
        return jsonify({"error": "No session_id"}), 400

    result = auto_enhance(sid)
    if result.get("error"):
        return jsonify({"error": result["error"]}), 500

    version, other_versions = _save_enhanced_to_db(sid)
    if not version:
        return jsonify({"error": "Failed to build enhanced rubric"}), 500

    return jsonify({"success": True, "version": version,
                     "other_versions": other_versions})


@app.route("/rubric-builder/refine/start", methods=["POST"])
def builder_refine_start():
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    body = request.get_json(silent=True) or {}
    sid = body.get("session_id")
    if not sid:
        return jsonify({"error": "No session_id"}), 400

    result = start_refinement(sid)
    if result.get("error"):
        return jsonify({"error": result["error"]}), 400
    return jsonify(result)


@app.route("/rubric-builder/refine/respond", methods=["POST"])
def builder_refine_respond():
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    body = request.get_json(silent=True) or {}
    sid = body.get("session_id")
    choice = body.get("choice")
    if not sid or not choice:
        return jsonify({"error": "Missing session_id or choice"}), 400

    result = respond_to_refinement(
        sid, choice,
        custom_points=body.get("custom_points"),
        custom_rationale=body.get("custom_rationale"),
    )
    if result.get("error"):
        return jsonify({"error": result["error"]}), 400
    return jsonify(result)


@app.route("/rubric-builder/refine/skip", methods=["POST"])
def builder_refine_skip():
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    body = request.get_json(silent=True) or {}
    sid = body.get("session_id")
    if not sid:
        return jsonify({"error": "No session_id"}), 400

    result = skip_question(sid)
    if result.get("error"):
        return jsonify({"error": result["error"]}), 400
    return jsonify({"phase": result["phase"], "current_q": result["current_q"],
                     "total_q": len(result["questions"])})


@app.route("/rubric-builder/refine/advance", methods=["POST"])
def builder_refine_advance():
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    body = request.get_json(silent=True) or {}
    sid = body.get("session_id")
    if not sid:
        return jsonify({"error": "No session_id"}), 400

    result = advance_question(sid)
    if result.get("error"):
        return jsonify({"error": result["error"]}), 400
    return jsonify({"phase": result["phase"], "current_q": result["current_q"],
                     "total_q": len(result["questions"])})


@app.route("/rubric-builder/refine/jump", methods=["POST"])
def builder_refine_jump():
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    body = request.get_json(silent=True) or {}
    sid = body.get("session_id")
    target_q = body.get("target_q")
    if not sid or target_q is None:
        return jsonify({"error": "Missing session_id or target_q"}), 400

    result = jump_to_question(sid, int(target_q))
    if result.get("error"):
        return jsonify({"error": result["error"]}), 400
    return jsonify({"phase": result["phase"], "current_q": result["current_q"],
                     "total_q": len(result["questions"])})


@app.route("/rubric-builder/save", methods=["POST"])
def builder_save():
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    body = request.get_json(silent=True) or {}
    sid = body.get("session_id")
    if not sid:
        return jsonify({"error": "No session_id"}), 400

    version, other_versions = _save_enhanced_to_db(sid)
    if not version:
        return jsonify({"error": "Session not found"}), 404

    flash(f"Enhanced rubric saved for Version {version}.", "success")
    return jsonify({"success": True, "version": version,
                     "other_versions": other_versions})


@app.route("/rubric-builder/session/<sid>")
def builder_session_status(sid):
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    sess = get_session(sid)
    if not sess:
        return jsonify({"error": "Session not found"}), 404
    return jsonify({
        "phase": sess["phase"],
        "current_q": sess["current_q"],
        "total_q": len(sess["questions"]),
        "questions": sess["questions"],
        "error": sess["error"],
    })


@app.route("/rubric-builder/session-enhanced/<version>")
def builder_view_enhanced(version):
    """Return the saved enhanced rubric for viewing/editing."""
    db = get_db()
    rubric = db.execute(
        "SELECT enhanced_rubric FROM rubrics WHERE version=? AND enhanced_rubric IS NOT NULL "
        "ORDER BY id DESC LIMIT 1", (version.upper(),)
    ).fetchone()
    if not rubric or not rubric["enhanced_rubric"]:
        return jsonify({"error": f"No enhanced rubric for version {version}"}), 404
    enhanced = json.loads(rubric["enhanced_rubric"])
    return jsonify(enhanced)


@app.route("/rubric-builder/update-enhanced", methods=["POST"])
def builder_update_enhanced():
    """Save inline edits to an enhanced rubric."""
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    body = request.get_json(silent=True)
    if not body or not body.get("version"):
        return jsonify({"error": "Missing data"}), 400

    version = body["version"].strip().upper()
    enhanced_json = json.dumps(body, indent=2)

    db = get_db()
    db.execute(
        "UPDATE rubrics SET enhanced_rubric=? WHERE version=? AND id = "
        "(SELECT MAX(id) FROM rubrics WHERE version=?)",
        (enhanced_json, version, version)
    )
    db.commit()
    return jsonify({"success": True, "version": version})


@app.route("/rubric-builder/remove-enhanced", methods=["POST"])
def builder_remove_enhanced():
    """Remove enhancement from a rubric version."""
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    body = request.get_json(silent=True) or {}
    version = body.get("version", "").strip().upper()
    if not version:
        return jsonify({"error": "No version specified"}), 400

    db = get_db()
    db.execute(
        "UPDATE rubrics SET enhanced_rubric=NULL WHERE version=?",
        (version,)
    )
    db.commit()
    return jsonify({"success": True, "version": version})


@app.route("/rubric-builder/map/generate", methods=["POST"])
def builder_map_generate():
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    body = request.get_json(silent=True) or {}
    primary_version = body.get("primary_version", "").strip().upper()
    target_version = body.get("target_version", "").strip().upper()
    if not primary_version or not target_version:
        return jsonify({"error": "Missing primary or target version"}), 400

    db = get_db()
    primary = db.execute(
        "SELECT enhanced_rubric FROM rubrics WHERE version=? AND enhanced_rubric IS NOT NULL "
        "ORDER BY id DESC LIMIT 1", (primary_version,)
    ).fetchone()
    if not primary or not primary["enhanced_rubric"]:
        return jsonify({"error": f"No enhanced rubric for version {primary_version}"}), 404

    target = db.execute(
        "SELECT content, total_points FROM rubrics WHERE version=? ORDER BY id DESC LIMIT 1",
        (target_version,)
    ).fetchone()
    if not target:
        return jsonify({"error": f"No rubric found for version {target_version}"}), 404

    enhanced = json.loads(primary["enhanced_rubric"])
    result = generate_mapping(enhanced, target["content"], target_version)
    if result.get("error"):
        return jsonify({"error": result["error"]}), 500

    return jsonify(result)


@app.route("/rubric-builder/map/apply", methods=["POST"])
def builder_map_apply():
    if not _HAS_BUILDER:
        return jsonify({"error": "Builder not available"}), 503
    body = request.get_json(silent=True) or {}
    primary_version = body.get("primary_version", "").strip().upper()
    target_version = body.get("target_version", "").strip().upper()
    mappings = body.get("mappings", [])
    if not primary_version or not target_version or not mappings:
        return jsonify({"error": "Missing required fields"}), 400

    db = get_db()
    primary = db.execute(
        "SELECT enhanced_rubric FROM rubrics WHERE version=? AND enhanced_rubric IS NOT NULL "
        "ORDER BY id DESC LIMIT 1", (primary_version,)
    ).fetchone()
    if not primary:
        return jsonify({"error": f"No enhanced rubric for version {primary_version}"}), 404

    target = db.execute(
        "SELECT total_points FROM rubrics WHERE version=? ORDER BY id DESC LIMIT 1",
        (target_version,)
    ).fetchone()
    if not target:
        return jsonify({"error": f"No rubric for version {target_version}"}), 404

    enhanced = json.loads(primary["enhanced_rubric"])
    mapped = apply_mapping(enhanced, mappings, target_version, target.get("total_points"))
    mapped_json = json.dumps(mapped, indent=2)

    db.execute(
        "UPDATE rubrics SET enhanced_rubric=? WHERE version=? AND id = "
        "(SELECT MAX(id) FROM rubrics WHERE version=?)",
        (mapped_json, target_version, target_version)
    )
    db.commit()

    return jsonify({"success": True, "version": target_version,
                     "questions_mapped": len(mapped["questions"])})


# ---------------------------------------------------------------------------
# Student Report PDF + Email Distribution
# ---------------------------------------------------------------------------

def _generate_student_pdf(exam_row, gd) -> bytes:
    """Generate a per-student grade report PDF (no scans) using reportlab."""
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.pagesizes import letter as rl_letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=rl_letter,
                            leftMargin=0.75*inch, rightMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)

    styles = getSampleStyleSheet()
    navy = rl_colors.HexColor("#1a2744")
    accent = rl_colors.HexColor("#3b82f6")
    green = rl_colors.HexColor("#16a34a")
    red_c = rl_colors.HexColor("#dc2626")
    light_bg = rl_colors.HexColor("#f8fafc")
    border_c = rl_colors.HexColor("#cbd5e1")

    title_style = ParagraphStyle("RPTitle", parent=styles["Title"],
                                 textColor=navy, fontSize=20, spaceAfter=4)
    subtitle_style = ParagraphStyle("RPSub", parent=styles["Normal"],
                                    textColor=rl_colors.HexColor("#475569"),
                                    fontSize=10, spaceAfter=12)
    heading_style = ParagraphStyle("RPH2", parent=styles["Heading2"],
                                   textColor=navy, fontSize=13, spaceBefore=16, spaceAfter=8)
    body_style = ParagraphStyle("RPBody", parent=styles["Normal"], fontSize=10, leading=14)
    fb_style = ParagraphStyle("RPFB", parent=styles["Normal"], fontSize=9, leading=12,
                              textColor=rl_colors.HexColor("#374151"))

    story = []

    # Header
    story.append(Paragraph("Exam Grade Report", title_style))
    total_possible = float(gd.get("total_possible", 0))
    total_earned = float(gd.get("total_earned", 0))
    pct = total_earned / total_possible * 100 if total_possible else 0
    grade = gd.get("letter_grade", "?")

    info_lines = [f"<b>{exam_row['student_name']}</b>"]
    if exam_row.get("student_sid"):
        info_lines.append(f"SID: {exam_row['student_sid']}")
    info_lines.append(f"Version {exam_row['version']} - Batch {exam_row['batch']}")
    story.append(Paragraph(" &nbsp;|&nbsp; ".join(info_lines), subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=border_c, spaceAfter=12))

    # Score summary
    grade_color = green if pct >= 70 else (rl_colors.HexColor("#d97706") if pct >= 60 else red_c)
    story.append(Paragraph(
        f"<b>Score:</b> {total_earned} / {total_possible} &nbsp;&nbsp; "
        f"<b>Percentage:</b> {pct:.1f}% &nbsp;&nbsp; "
        f"<font color='{grade_color.hexval()}'><b>Grade: {grade}</b></font>",
        body_style
    ))
    story.append(Spacer(1, 12))

    # Question breakdown table
    scores = gd.get("scores", [])
    if scores:
        story.append(Paragraph("Question Breakdown", heading_style))
        header = [
            Paragraph("<b>Question</b>", fb_style),
            Paragraph("<b>Earned</b>", fb_style),
            Paragraph("<b>Max</b>", fb_style),
            Paragraph("<b>Feedback</b>", fb_style),
        ]
        table_data = [header]
        for s in scores:
            fb_text = str(s.get("feedback", "")).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            table_data.append([
                Paragraph(str(s["question"]), fb_style),
                Paragraph(str(s["earned_points"]), fb_style),
                Paragraph(str(s["max_points"]), fb_style),
                Paragraph(fb_text, fb_style),
            ])

        col_widths = [0.7*inch, 0.7*inch, 0.6*inch, 4.5*inch]
        t = Table(table_data, colWidths=col_widths, repeatRows=1)

        # Row-level shading
        t_style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#e2e8f0")),
            ("TEXTCOLOR", (0, 0), (-1, 0), navy),
            ("GRID", (0, 0), (-1, -1), 0.5, border_c),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ]
        for idx, s in enumerate(scores, start=1):
            ep = float(s.get("earned_points", 0))
            mp = float(s.get("max_points", 1))
            if ep == mp:
                t_style_cmds.append(("BACKGROUND", (0, idx), (-1, idx), rl_colors.HexColor("#f0fdf4")))
            elif mp > 0 and ep / mp < 0.7:
                t_style_cmds.append(("BACKGROUND", (0, idx), (-1, idx), rl_colors.HexColor("#fef2f2")))
            else:
                t_style_cmds.append(("BACKGROUND", (0, idx), (-1, idx), light_bg))

        t.setStyle(TableStyle(t_style_cmds))
        story.append(t)

    # Overall feedback
    overall = gd.get("overall_feedback", "")
    if overall:
        story.append(Paragraph("Overall Feedback", heading_style))
        overall_escaped = overall.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(overall_escaped, body_style))

    doc.build(story)
    return buf.getvalue()


try:
    from gmail_sender import send_email as gmail_send, is_configured as gmail_is_configured, get_sender_email
    _HAS_GMAIL = True
except ImportError:
    _HAS_GMAIL = False


@app.route("/email/preview", methods=["POST"])
def email_preview():
    """Return count of students with email addresses for the current filter."""
    db = get_db()
    version_filter = request.json.get("version", "").strip().upper() or None
    batch_filter = request.json.get("batch")
    try:
        batch_filter = int(batch_filter)
    except (ValueError, TypeError):
        batch_filter = None

    query = """
        SELECT COUNT(*) FROM exams e
        JOIN roster r ON e.student_sid = r.sid
        WHERE e.grade_data IS NOT NULL AND e.reviewed = 1
        AND r.email IS NOT NULL AND r.email != ''
    """
    params = []
    if version_filter:
        query += " AND e.version = ?"
        params.append(version_filter)
    if batch_filter:
        query += " AND e.batch = ?"
        params.append(batch_filter)

    count = db.execute(query, params).fetchone()[0]
    total_graded = db.execute(
        "SELECT COUNT(*) FROM exams WHERE grade_data IS NOT NULL"
    ).fetchone()[0]
    total_reviewed = db.execute(
        "SELECT COUNT(*) FROM exams WHERE grade_data IS NOT NULL AND reviewed = 1"
    ).fetchone()[0]

    gmail_ok = _HAS_GMAIL and gmail_is_configured()
    sender = get_sender_email() if gmail_ok else None

    return jsonify({
        "email_ready": count,
        "total_graded": total_graded,
        "total_reviewed": total_reviewed,
        "gmail_configured": gmail_ok,
        "sender_email": sender,
    })


@app.route("/email/send", methods=["POST"])
def email_send():
    """Send grade report PDFs to students via Gmail OAuth2."""
    if not _HAS_GMAIL or not gmail_is_configured():
        return jsonify({"error": "Gmail not configured. Run: python gmail_sender.py --setup berkeley"}), 500

    db = get_db()
    body = request.get_json()
    if not body:
        return jsonify({"error": "Missing JSON body"}), 400

    subject = body.get("subject", "").strip()
    message = body.get("message", "").strip()
    test_mode = body.get("test_mode", True)
    test_email = body.get("test_email", "").strip()
    test_limit = int(body.get("test_limit", 10))
    version_filter = body.get("version", "").strip().upper() or None
    batch_filter = body.get("batch")
    try:
        batch_filter = int(batch_filter)
    except (ValueError, TypeError):
        batch_filter = None

    if not subject or not message:
        return jsonify({"error": "Subject and message are required"}), 400
    if test_mode and not test_email:
        return jsonify({"error": "Test email address is required in test mode"}), 400

    # Query exams with roster email
    query = """
        SELECT e.*, r.email, r.first_name, r.last_name FROM exams e
        JOIN roster r ON e.student_sid = r.sid
        WHERE e.grade_data IS NOT NULL AND e.reviewed = 1
        AND r.email IS NOT NULL AND r.email != ''
    """
    params = []
    if version_filter:
        query += " AND e.version = ?"
        params.append(version_filter)
    if batch_filter:
        query += " AND e.batch = ?"
        params.append(batch_filter)
    query += " ORDER BY e.student_name"

    rows = db.execute(query, params).fetchall()
    if not rows:
        return jsonify({"error": "No students with email addresses found"}), 404

    if test_mode:
        rows = rows[:test_limit]

    sent = 0
    errors = []
    for row in rows:
        try:
            gd = json.loads(row["grade_data"])
            pdf_bytes = _generate_student_pdf(dict(row), gd)

            personal_msg = message.replace("{name}", row["first_name"])
            recipient = test_email if test_mode else row["email"]
            safe_name = re.sub(r'[^\w\-.]', '_', row["student_name"])

            gmail_send(
                to=recipient,
                subject=subject,
                body=personal_msg,
                pdf_bytes=pdf_bytes,
                pdf_filename=f"Grade_Report_{safe_name}.pdf",
            )
            sent += 1
        except Exception as e:
            errors.append({"student": row["student_name"], "error": str(e)})

    result = {"sent": sent, "total": len(rows)}
    if test_mode:
        result["test_email"] = test_email
    if errors:
        result["errors"] = errors
    return jsonify(result)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    init_db()
    port = int(os.environ.get("PORT", "5000"))
    print("\n Exam Grader running locally.")
    print("   Cover-page OCR uses local Ollama model — no student data leaves this machine.")
    print(f"   OCR model : {OLLAMA_VISION_MODEL}  (change OLLAMA_VISION_MODEL to swap)")
    print("   Grading   : Claude Sonnet via Anthropic API (anonymous IDs only)")
    print(f"   Open: http://localhost:{port}\n")
    app.run(debug=os.environ.get("FLASK_DEBUG", "").lower() == "true",
            host="127.0.0.1", port=port, use_reloader=False)
